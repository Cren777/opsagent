from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "build" / "midterm_acceptance"
FIG = OUT / "figures"

BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C3D0"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        try:
            return ImageFont.truetype(item, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for para in text.split("\n"):
        current = ""
        for ch in para:
            trial = current + ch
            if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines or [""]


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: str = "",
    fill: str = "#FFFFFF",
    outline: str = "#7B8EA8",
):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    tf = font(30, True)
    bf = font(22)
    draw.text((x1 + 22, y1 + 16), title, fill="#17365D", font=tf)
    if body:
        y = y1 + 58
        for line in wrap_text(draw, body, bf, x2 - x1 - 44):
            draw.text((x1 + 22, y), line, fill="#263238", font=bf)
            y += 30


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#406C99"):
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - direction * 18, ey - 10), (ex - direction * 18, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 10, ey - direction * 18), (ex + 10, ey - direction * 18)]
    draw.polygon(pts, fill=color)


def save_fig_architecture():
    img = Image.new("RGB", (1500, 900), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title_f = font(42, True)
    d.text((55, 35), "OpsAgent 系统总体架构", fill="#0B2545", font=title_f)
    draw_box(d, (80, 140, 360, 280), "用户层", "Web 聊天界面\n知识库/日志/配置管理", "#F7FAFC")
    draw_box(d, (520, 110, 940, 310), "应用服务层", "FastAPI + API 路由\nSSE 流式输出\n静态 SPA 托管", "#E8EEF5")
    draw_box(d, (1110, 140, 1430, 280), "前端层", "Vue 3 + TypeScript\nElement Plus + Pinia", "#F7FAFC")
    draw_box(d, (140, 420, 420, 610), "核心编排", "意图分类\n任务路由\n结果融合", "#FFF8E1", "#C9A227")
    draw_box(d, (550, 400, 850, 630), "能力模块", "RAG 知识检索\nText2SQL 数据分析\n日志故障排查\n诊断脚本执行", "#EEF7ED", "#5B8A54")
    draw_box(d, (1030, 400, 1360, 630), "模型与工具", "DeepSeek / DashScope\nBGE Embedding\nMySQL / ClickHouse / CSV", "#F2F4F7")
    draw_box(d, (260, 730, 1260, 850), "数据存储层", "SQLite 配置库、Milvus Lite 向量库、知识库文档、上传日志、故障案例库、脚本白名单", "#F7FAFC")
    arrow(d, (360, 210), (520, 210))
    arrow(d, (940, 210), (1110, 210))
    arrow(d, (700, 310), (700, 400))
    arrow(d, (420, 515), (550, 515))
    arrow(d, (850, 515), (1030, 515))
    arrow(d, (700, 630), (700, 730))
    img.save(FIG / "architecture.png", quality=95)


def save_fig_flow():
    img = Image.new("RGB", (1500, 900), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((55, 35), "核心业务流程", fill="#0B2545", font=font(42, True))
    boxes = [
        ((60, 150, 300, 270), "提交问题", "自然语言提问\n可附加日志"),
        ((420, 150, 670, 270), "意图识别", "规则快速匹配\nLLM 精准分类"),
        ((790, 80, 1110, 205), "知识问答", "Milvus 检索\nLLM 融合回答"),
        ((790, 260, 1110, 385), "数据分析", "Schema 探查\n生成/校验/执行 SQL"),
        ((790, 440, 1110, 565), "故障排查", "日志上下文\n诊断脚本\n历史案例"),
        ((1230, 260, 1460, 385), "返回结果", "答案、来源\nSQL、诊断元数据"),
        ((520, 680, 1180, 820), "沉淀复用", "知识文档索引、日志索引、故障案例自动保存，供后续相似问题复用"),
    ]
    for xy, title, body in boxes:
        draw_box(d, xy, title, body, "#F7FAFC")
    arrow(d, (300, 210), (420, 210))
    arrow(d, (670, 210), (790, 145))
    arrow(d, (670, 210), (790, 322))
    arrow(d, (670, 210), (790, 502))
    arrow(d, (1110, 145), (1230, 322))
    arrow(d, (1110, 322), (1230, 322))
    arrow(d, (1110, 502), (1230, 322))
    arrow(d, (920, 565), (920, 680))
    img.save(FIG / "flow.png", quality=95)


def save_fig_modules():
    img = Image.new("RGB", (1500, 900), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((55, 35), "功能模块划分", fill="#0B2545", font=font(42, True))
    modules = [
        ((60, 130, 410, 300), "智能对话", "非流式/流式聊天\n多意图处理\n上下文历史"),
        ((535, 130, 885, 300), "知识库管理", "上传 md/txt\n目录分类\n预览与重建索引"),
        ((1010, 130, 1450, 300), "数据源配置", "MySQL / ClickHouse\nExcel / CSV\n连接测试与激活"),
        ((60, 390, 410, 560), "日志与案例", "日志上传/脱敏\n分类筛选\n案例沉淀复用"),
        ((535, 390, 885, 560), "诊断工具", "脚本白名单\n待启用审核\n执行与输出截断"),
        ((1010, 390, 1450, 560), "索引管理", "知识库/日志/案例\ncollection 状态\n重建与清理"),
        ((300, 675, 700, 820), "大模型配置", "OpenAI 兼容接口\nDashScope\n主力模型切换"),
        ((850, 675, 1250, 820), "安全与审计", "SQL 只读校验\n路径限制\n密钥加密与脱敏"),
    ]
    for xy, title, body in modules:
        draw_box(d, xy, title, body, "#F7FAFC")
    img.save(FIG / "modules.png", quality=95)


def save_fig_deployment():
    img = Image.new("RGB", (1500, 850), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((55, 35), "部署与运行视图", fill="#0B2545", font=font(42, True))
    draw_box(d, (90, 140, 420, 300), "浏览器", "访问 Vue SPA\n调用 /api/*", "#F7FAFC")
    draw_box(d, (585, 110, 940, 330), "FastAPI 服务", "uvicorn\n8080 端口\nCORS / API Key", "#E8EEF5")
    draw_box(d, (1080, 140, 1410, 300), "外部模型", "DeepSeek API\n阿里云百炼 DashScope", "#F7FAFC")
    draw_box(d, (150, 510, 470, 700), "本地数据", "SQLite 配置库\n知识/日志/案例文件\nMilvus Lite", "#EEF7ED", "#5B8A54")
    draw_box(d, (605, 510, 925, 700), "业务数据源", "MySQL\nClickHouse\nExcel/CSV", "#FFF8E1", "#C9A227")
    draw_box(d, (1060, 510, 1380, 700), "诊断脚本", "scripts/approved\n30 秒超时\n输出截断", "#F2F4F7")
    arrow(d, (420, 220), (585, 220))
    arrow(d, (940, 220), (1080, 220))
    arrow(d, (760, 330), (310, 510))
    arrow(d, (760, 330), (765, 510))
    arrow(d, (760, 330), (1220, 510))
    img.save(FIG / "deployment.png", quality=95)


def save_fig_data_model():
    img = Image.new("RGB", (1500, 900), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((55, 35), "核心数据对象关系", fill="#0B2545", font=font(42, True))
    boxes = [
        ((80, 150, 430, 320), "LLMProvider", "id, name, provider_type\napi_key(加密), model\nis_primary"),
        ((560, 150, 910, 320), "DataSource", "id, type, host/path\ncredentials(加密)\nis_active"),
        ((1040, 150, 1390, 320), "KnowledgeFile", "file_id, path, folder\nmtime, index_status"),
        ((170, 500, 520, 690), "LogMetadata", "file_id, filename, source\ncategory, severity, tags\nanalysis, stored_path"),
        ((740, 500, 1090, 690), "IncidentCase", "case_id, symptoms\nroot_cause, solution\nstatus, category, evidence"),
    ]
    for xy, title, body in boxes:
        draw_box(d, xy, title, body, "#F7FAFC")
    arrow(d, (735, 320), (345, 500))
    arrow(d, (345, 690), (740, 595))
    arrow(d, (1215, 320), (915, 500))
    d.text((95, 770), "说明：配置数据存入 SQLite，知识/日志/案例内容进入 Milvus Lite 后参与相似检索；敏感字段加密或脱敏后展示。", fill="#263238", font=font(26))
    img.save(FIG / "data_model.png", quality=95)


def build_figures():
    FIG.mkdir(parents=True, exist_ok=True)
    save_fig_architecture()
    save_fig_flow()
    save_fig_modules()
    save_fig_deployment()
    save_fig_data_model()


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths: list[float]):
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        cell.text = text
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = "Calibri"
                set_east_asia(run, "微软雅黑")
                run.font.size = Pt(9.5)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    set_east_asia(run, "宋体")
                    run.font.size = Pt(9)
    if widths:
        set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def set_doc_defaults(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def set_east_asia(run, name: str):
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)


def add_header_footer(doc: Document, label: str):
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = label
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)
        set_east_asia(run, "微软雅黑")
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("第 ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    footer_p.add_run(" 页")


def add_cover(doc: Document, title: str, subtitle: str, kind: str):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    set_east_asia(run, "微软雅黑")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE
    set_east_asia(run, "微软雅黑")

    doc.add_paragraph()
    meta = [
        ["项目名称", "OpsAgent 智能运维助手"],
        ["文档类型", kind],
        ["阶段", "中期验收"],
        ["版本", "V1.0"],
        ["日期", "2026 年 6 月"],
    ]
    add_table(doc, ["项目", "内容"], meta, [1.5, 4.8])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("说明：本文档为 Word 版中期验收文档，按最终验收文档格式准备，可在后续阶段继续补充完善。")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(85, 85, 85)
    set_east_asia(run, "宋体")
    doc.add_page_break()


def add_toc_placeholder(doc: Document, entries: list[str]):
    doc.add_heading("目录", level=1)
    p = doc.add_paragraph()
    p.add_run("提示：在 Word 中打开后，可右键更新目录域，生成正式页码目录。")
    for item in entries:
        doc.add_paragraph(item)
    doc.add_page_break()


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_figure(doc: Document, name: str, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FIG / name), width=Inches(6.35))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in c.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)


def add_requirements_expansion(doc: Document):
    doc.add_heading("8 详细业务需求分析", 1)
    doc.add_heading("8.1 系统边界", 2)
    add_table(
        doc,
        ["边界项", "纳入范围", "暂不纳入范围"],
        [
            ["运维知识问答", "知识文档检索、来源引用、步骤化回答、索引重建。", "知识内容自动采编、知识审核流和多人协同编辑。"],
            ["自然语言数据分析", "面向已配置数据源的 SELECT 查询、结果摘要、SQL 展示。", "自动写入数据库、复杂数据仓库建模、跨源联邦查询优化。"],
            ["日志故障排查", "日志上传、目录发现、脱敏预览、摘要分析、诊断上下文注入。", "生产日志实时采集 Agent、日志告警规则平台和分布式链路追踪。"],
            ["自动诊断脚本", "白名单脚本管理、待启用审核、安全执行和输出截断。", "任意 Shell 交互、远程批量执行和高危修复命令自动执行。"],
            ["案例复用", "自动保存故障案例、分类筛选、相似案例匹配。", "完整 ITSM 工单流转、审批、SLA 计算和通知系统。"],
        ],
        [1.25, 3.0, 2.25],
    )

    doc.add_heading("8.2 业务术语", 2)
    add_table(
        doc,
        ["术语", "定义"],
        [
            ["RAG", "检索增强生成。先从知识库或日志向量库检索相关证据，再交给 LLM 生成答案。"],
            ["Text2SQL", "将自然语言问题转换为 SQL 查询，并在安全校验后执行。"],
            ["意图", "系统对用户问题类型的分类，包括 knowledge_query、data_analysis、fault_troubleshooting。"],
            ["附件上下文", "用户上传日志或系统自动匹配日志后，提取出的错误数、警告数、关键片段等诊断证据。"],
            ["故障案例", "一次故障排查形成的症状、根因、处理方案和证据集合。"],
            ["索引状态", "知识、日志或案例向量索引是否已构建、是否因文件修改而待重建的状态。"],
        ],
        [1.35, 5.15],
    )

    doc.add_heading("8.3 详细功能需求清单", 2)
    add_table(
        doc,
        ["编号", "需求名称", "触发条件", "输入", "输出/结果", "验收方式"],
        [
            ["FR-01", "普通聊天入口", "用户在首页输入问题。", "query、history、datasource_id、attachments。", "返回 answer、intent、sources、sql、diagnostics。", "调用 /api/chat 返回结构完整。"],
            ["FR-02", "SSE 流式回答", "用户发送聊天并启用流式接口。", "ChatRequest。", "intent 事件、多个 token 事件、done 事件。", "浏览器端逐块显示回答。"],
            ["FR-03", "日志附件强制故障排查", "attachments 中存在 type=log。", "日志 file_id。", "intent 被解析为 fault_troubleshooting。", "单元测试和上传日志演示。"],
            ["FR-04", "文本中日志文件名解析", "问题中出现 *.log/*.txt/*.out/*.gz。", "用户自然语言。", "自动匹配日志目录中的文件作为附件。", "指定 ops_agent_*.log 能进入诊断。"],
            ["FR-05", "知识库文件上传", "管理员上传 .md/.txt。", "文件、目录。", "文件保存并出现在目录树。", "页面上传后可预览。"],
            ["FR-06", "知识库目录管理", "管理员创建/重命名/删除目录。", "folder path/name。", "目录树更新，文件过滤正确。", "页面操作验证。"],
            ["FR-07", "知识检索问答", "用户询问操作方法。", "query。", "返回操作步骤与来源文件。", "示例问题返回知识库来源。"],
            ["FR-08", "数据源 CRUD", "管理员维护数据源。", "MySQL/ClickHouse/CSV 配置。", "创建、更新、删除、激活数据源。", "配置页面和接口联调。"],
            ["FR-09", "数据源连接测试", "保存前或保存后点击测试。", "连接参数。", "返回连接成功/失败及错误信息。", "使用测试按钮验证。"],
            ["FR-10", "自然语言生成 SQL", "用户提出统计查询。", "query、schema。", "生成 SELECT SQL。", "SQL 包含 LIMIT 且可执行。"],
            ["FR-11", "SQL 安全拦截", "模型生成危险 SQL。", "SQL 字符串。", "拦截写操作、超大 LIMIT、注入模式。", "pytest 覆盖危险关键字。"],
            ["FR-12", "日志上传", "用户上传日志。", "filename、content、category。", "metadata、分析摘要、severity。", "上传后列表可见。"],
            ["FR-13", "日志脱敏预览", "用户查看日志详情。", "file_id。", "敏感字段替换后内容。", "password/token 不明文展示。"],
            ["FR-14", "日志分类筛选", "用户按分类、来源、严重级别筛选。", "query、category、source、severity。", "返回匹配日志列表。", "管理页面筛选验证。"],
            ["FR-15", "故障案例自动保存", "故障排查完成且未命中案例。", "query、answer、symptoms、evidence。", "新增 case 记录。", "案例列表出现新记录。"],
            ["FR-16", "相似案例复用", "新问题与历史案例相似。", "query、symptoms。", "返回历史根因和方案。", "构造相似问题验证命中。"],
            ["FR-17", "案例状态管理", "管理员标记案例状态。", "case_id、status。", "状态更新。", "页面更新后列表刷新。"],
            ["FR-18", "诊断脚本上传审核", "用户上传脚本。", "check_*.sh 或 check_*.py。", "进入 pending，不可直接执行。", "pending 列表可见。"],
            ["FR-19", "白名单脚本执行", "用户选择已启用脚本。", "script_name、args。", "stdout、stderr、exit_code。", "approved 目录脚本可运行。"],
            ["FR-20", "索引重建", "知识/日志/案例变更后。", "collection 类型。", "重建完成并更新状态。", "索引页面操作验证。"],
        ],
        [0.55, 1.15, 1.25, 1.25, 1.45, 1.1],
    )

    doc.add_heading("9 用例规约", 1)
    use_cases = [
        ["UC-01", "知识问答", "一线运维人员", "系统已启动，知识库索引可用。", "输入“如何排查 Linux 磁盘满”。", "返回步骤、命令和来源文件。", "知识库无命中时提示补充知识文档。"],
        ["UC-02", "自然语言查数", "数据分析人员", "已配置并激活数据源。", "输入“最近 7 天告警数量按级别统计”。", "生成 SELECT SQL，执行后给出摘要。", "数据源不可用时提示去配置页面检查。"],
        ["UC-03", "上传日志诊断", "一线运维人员", "日志文件扩展名合法且小于 20MB。", "上传日志并输入“帮我分析这个日志”。", "强制进入故障排查，返回错误摘要和处理建议。", "日志过大或扩展名非法时拒绝上传。"],
        ["UC-04", "指定已有日志诊断", "一线运维人员", "logs 或 data/logs 存在对应文件。", "输入“分析 ops_agent_2026-05-25.log”。", "系统自动解析文件名并注入日志上下文。", "文件不存在时按普通问题处理或提示未匹配。"],
        ["UC-05", "管理知识文档", "运维负责人", "拥有管理页面访问权限。", "创建目录、上传 md/txt、预览、重建索引。", "文档出现在树中，索引状态更新。", "路径非法或扩展名非法时拒绝。"],
        ["UC-06", "维护故障案例", "运维负责人", "案例库已有记录。", "按分类筛选，修改状态或分类。", "列表和详情同步更新。", "案例不存在时返回 404 或失败提示。"],
        ["UC-07", "启用诊断脚本", "系统管理员", "pending 目录存在待审核脚本。", "预览脚本后点击启用。", "脚本移动到 approved 并可执行。", "脚本名不符合 check_*.sh/py 时拒绝。"],
        ["UC-08", "配置模型提供商", "系统管理员", "具备 API Key。", "新增 Provider，测试并设为主力。", "聊天接口使用配置库中的主力模型。", "测试失败时不影响原主力模型。"],
    ]
    add_table(doc, ["用例编号", "用例名称", "参与者", "前置条件", "主事件流", "成功后置条件", "异常流"], use_cases, [0.65, 1.0, 1.0, 1.35, 1.6, 1.45, 1.5])

    doc.add_heading("10 非功能需求量化指标", 1)
    add_table(
        doc,
        ["类别", "指标", "中期目标", "最终目标"],
        [
            ["性能", "普通管理接口响应时间", "本地环境多数接口 2 秒内返回。", "典型接口 1 秒内返回，索引重建类接口异步提示。"],
            ["性能", "聊天首包时间", "流式接口能先返回 intent 事件。", "模型可用时 3 秒内开始输出 token。"],
            ["安全", "SQL 风险控制", "危险关键字和 LIMIT 约束已有测试。", "补充更多注入样例和审计日志。"],
            ["可靠性", "模型调用失败", "LLM 分类失败回退规则分类。", "Provider 级自动回退并记录错误原因。"],
            ["可维护性", "模块边界", "已按路由、服务、模型、核心编排拆分。", "补充接口文档、用户手册和部署脚本。"],
            ["可用性", "管理页面", "核心 CRUD 页面可演示。", "补充空状态、加载态、批量操作和错误提示。"],
        ],
        [1.0, 1.55, 2.2, 2.75],
    )

    doc.add_heading("11 中期验收材料清单", 1)
    add_table(
        doc,
        ["材料", "当前交付内容", "验收时使用方式"],
        [
            ["需求分析说明书", "项目背景、开题意见回应、角色场景、功能/非功能需求、用例、演示清单。", "说明系统要解决什么问题、做了哪些功能和如何验收。"],
            ["概要设计说明书", "总体架构、模块划分、接口分组、部署运行、安全概要和测试策略。", "说明系统整体怎么拆、模块之间怎么协作。"],
            ["详细设计说明书", "包结构、核心类、接口、数据对象、算法、安全边界、测试计划。", "说明核心功能如何落到代码和数据结构。"],
            ["功能演示", "智能对话、知识库、数据源、日志案例、诊断工具、索引管理、模型配置页面。", "按演示脚本逐项展示完成进度。"],
            ["测试记录", "pytest 测试、前端类型检查命令、后端语法检查命令。", "最终验收前补充截图和日志。"],
        ],
        [1.1, 3.0, 2.4],
    )


def add_outline_expansion(doc: Document):
    doc.add_heading("9 架构设计决策", 1)
    add_table(
        doc,
        ["决策项", "选择", "原因", "影响"],
        [
            ["前后端架构", "Vue SPA + FastAPI API", "前端交互复杂、后端需要统一编排模型和数据能力。", "开发阶段可独立运行，构建后可由 FastAPI 托管。"],
            ["流式输出", "SSE", "聊天场景需要低成本单向流式文本，浏览器原生支持较好。", "接口返回 intent/token/done，前端逐块渲染。"],
            ["向量库", "Milvus Lite", "中期项目便于本地运行，不依赖独立服务。", "适合演示和小规模知识库，最终可迁移到服务版 Milvus。"],
            ["配置库", "SQLite + SQLAlchemy + Fernet", "本地部署简单，同时需要保存数据源和模型密钥。", "敏感字段加密，数据库文件不提交源码。"],
            ["SQL 安全", "生成后强校验", "LLM 生成不可完全信任，必须在执行前约束。", "只读 SELECT 和 LIMIT 上限成为硬边界。"],
            ["脚本执行", "白名单目录", "自动化诊断有执行风险，需要强限制。", "只执行 approved 脚本，上传脚本需审核启用。"],
        ],
        [1.0, 1.35, 2.35, 2.0],
    )

    doc.add_heading("10 模块协作关系", 1)
    add_table(
        doc,
        ["发起模块", "协作模块", "调用关系", "数据传递"],
        [
            ["ChatView", "chat API", "提交 ChatRequest 或建立 SSE 请求。", "query、history、datasource_id、attachments。"],
            ["chat API", "Orchestrator", "调用 process/process_stream。", "请求对象转为编排参数。"],
            ["Orchestrator", "IntentClassifier", "先分类再路由。", "query -> IntentResult。"],
            ["Orchestrator", "LogUploadService", "解析提到的日志并生成附件上下文。", "filename/file_id -> analysis context。"],
            ["Orchestrator", "Text2SQLGenerator", "数据分析意图下生成 SQL。", "query + schema -> sql。"],
            ["Text2SQLGenerator", "DataSource", "SQL 校验后执行。", "sql -> rows。"],
            ["Orchestrator", "IncidentCaseMemory", "故障排查前查相似案例，结束后保存案例。", "symptoms/root_cause/solution/evidence。"],
            ["IndexManagementView", "IndexService", "触发索引状态查询或重建。", "collection -> status/rebuild result。"],
        ],
        [1.1, 1.35, 2.0, 3.05],
    )

    doc.add_heading("11 运行时数据与目录规划", 1)
    add_table(
        doc,
        ["路径/资源", "用途", "管理策略"],
        [
            ["data/app_config.db", "保存数据源和 LLM Provider 配置。", "运行产物，不提交；敏感字段加密。"],
            ["data/incident_cases.db", "保存故障案例。", "运行产物，不提交；可通过案例页面管理。"],
            ["data/knowledge", "保存知识库 md/txt 文件。", "源码可带样例，业务文档通过页面维护。"],
            ["data/uploads/logs", "保存用户上传日志和 metadata。", "运行产物，不提交；预览时脱敏。"],
            ["logs", "保存系统运行日志。", "可被日志目录服务发现并用于诊断。"],
            ["data/vectors/milvus.db", "Milvus Lite 向量索引文件。", "运行产物，不提交；索引页面可重建。"],
            ["scripts/approved", "可执行诊断脚本白名单。", "只允许 check_*.sh/py 等安全脚本。"],
            ["scripts/pending", "待审核上传脚本。", "默认不可执行，启用后才进入 approved。"],
        ],
        [1.7, 2.35, 2.95],
    )

    doc.add_heading("12 接口矩阵", 1)
    add_table(
        doc,
        ["分组", "路径", "方法", "请求要点", "响应要点"],
        [
            ["聊天", "/api/chat", "POST", "query、history、datasource_id、attachments。", "answer、intent、sources、sql、diagnostics。"],
            ["聊天", "/api/chat/stream", "POST", "同 /api/chat。", "SSE: intent/token/done/error。"],
            ["数据源", "/api/config/datasources", "GET/POST", "创建时传 type 和 config。", "数据源列表或新建结果。"],
            ["数据源", "/api/config/datasources/{id}/test", "POST", "已保存数据源 id。", "连接测试结果。"],
            ["数据源", "/api/config/datasources/{id}/tables", "GET", "数据源 id。", "表名、字段和样例结构。"],
            ["模型", "/api/config/llm", "GET/POST", "provider_type、api_key、base_url、model。", "Provider 列表或新建结果。"],
            ["模型", "/api/config/llm/{id}/primary", "POST", "Provider id。", "主力模型切换结果。"],
            ["知识库", "/api/knowledge/tree", "GET", "无。", "目录树和文件节点。"],
            ["知识库", "/api/knowledge/upload", "POST", "文件和目标目录。", "文件 metadata。"],
            ["日志", "/api/uploads/logs", "POST", "上传日志文件和分类。", "日志 metadata 和分析摘要。"],
            ["日志案例", "/api/incidents/logs/{file_id}/preview", "GET", "日志 file_id。", "脱敏后的内容和 metadata。"],
            ["日志案例", "/api/incidents/incidents/{case_id}/status", "PUT", "case_id、status。", "状态更新结果。"],
            ["诊断", "/api/diagnostics/scripts/{script_name}/run", "POST", "脚本名和参数。", "stdout、stderr、exit_code。"],
            ["索引", "/api/indexes/{collection}/clear", "POST", "collection 名。", "清理结果。"],
        ],
        [0.85, 1.9, 0.75, 2.0, 2.0],
    )

    doc.add_heading("13 关键流程概要", 1)
    doc.add_heading("13.1 故障排查流程", 2)
    add_numbered(
        doc,
        [
            "用户上传日志或在问题中提到日志文件名。",
            "LogUploadService 保存或发现日志，并生成错误数、警告数、关键模式、摘要片段。",
            "Orchestrator 将日志附件合并到请求上下文，并将意图修正为 fault_troubleshooting。",
            "系统先查询 IncidentCaseMemory，命中高相似案例时直接返回历史根因和解决方案。",
            "未命中时并行获取知识库上下文、日志索引结果和诊断脚本输出。",
            "ResponseFusion 将多源证据整理为故障现象、可能原因、排查步骤和处理建议。",
            "处理结果自动保存为故障案例，进入后续复用链路。",
        ],
    )
    doc.add_heading("13.2 Text2SQL 流程", 2)
    add_numbered(
        doc,
        [
            "用户提出数据分析问题。",
            "系统获取活跃数据源，并通过 SchemaManager 获取表结构。",
            "LLM 根据自然语言和 schema 生成候选 SQL。",
            "SQLValidator 执行只读、LIMIT 和注入模式校验。",
            "数据源执行 SQL 并返回 rows。",
            "ResponseFusion 将查询结果转为自然语言摘要，同时在响应元数据中保留 SQL。",
        ],
    )

    doc.add_heading("14 概要层风险与对策", 1)
    add_table(
        doc,
        ["风险", "影响", "中期对策", "最终阶段计划"],
        [
            ["LLM 输出不稳定", "意图或 SQL 生成可能不准确。", "规则分类兜底，SQL 执行前强校验。", "增加提示词模板和失败重试策略。"],
            ["向量检索召回不足", "知识问答可能缺少证据。", "按知识、日志、案例分别维护 collection。", "增加 chunk 策略调优和召回评估集。"],
            ["日志内容敏感", "预览或上下文可能泄露密钥。", "预览和摘要均脱敏。", "增加更多敏感字段模式和审计。"],
            ["脚本执行风险", "误执行高危命令。", "白名单、文件名限制、待启用审核。", "增加脚本元数据、权限和执行审计。"],
            ["演示环境依赖", "模型、数据库或 embedding 依赖不可用会影响演示。", "提供样例数据和错误提示。", "准备离线演示脚本和截图备份。"],
        ],
        [1.35, 1.45, 2.25, 2.45],
    )


def add_detail_expansion(doc: Document):
    doc.add_heading("13 核心类与方法设计", 1)
    add_table(
        doc,
        ["类/服务", "主要方法", "输入", "输出", "设计要点"],
        [
            ["Orchestrator", "process、process_stream", "query、datasource_id、history、attachments", "result 或 SSE event", "统一编排入口，合并日志附件，修正意图并路由处理器。"],
            ["IntentClassifier", "classify、_rule_classify、_llm_classify", "query", "IntentResult", "规则高置信命中直接返回，低置信再调用 LLM。"],
            ["TaskRouter", "register、route", "IntentType、handler、kwargs", "处理器结果", "将意图与处理函数解耦，便于扩展新任务。"],
            ["ResponseFusion", "fuse_for_knowledge、fuse_for_data、fuse_for_troubleshooting", "query、context/results", "answer", "统一回答风格，把证据转为结构化文本。"],
            ["Text2SQLGenerator", "generate", "自然语言和 schema", "SQL", "生成后交由 SQLValidator 检查。"],
            ["SQLValidator", "validate", "SQL 字符串", "异常或通过", "只读 SELECT、必须 LIMIT、禁止写操作和注入片段。"],
            ["LogUploadService", "save_log_file、list_logs、preview_log、resolve_mentioned_logs", "文件、file_id、query", "metadata、content、attachments", "统一上传日志和本地日志发现。"],
            ["IncidentCaseMemory", "save_case、find_similar、list_cases", "query、symptoms、case_id", "case 或列表", "SQLite 保存案例，基于 token overlap 匹配相似问题。"],
            ["DiagnosticService", "upload_script、enable_script、run_script", "script_name、content、args", "metadata 或执行输出", "待启用区和 approved 区分离。"],
        ],
        [1.35, 1.55, 1.55, 1.15, 1.95],
    )

    doc.add_heading("14 请求与响应结构设计", 1)
    add_table(
        doc,
        ["对象", "字段", "字段说明", "约束"],
        [
            ["ChatRequest", "query", "用户自然语言问题。", "必填，字符串。"],
            ["ChatRequest", "history", "历史消息数组，用于多轮上下文。", "可为空，元素为 dict。"],
            ["ChatRequest", "datasource_id", "临时指定数据源。", "可为空。"],
            ["ChatRequest", "attachments", "附件数组，日志附件 type=log。", "可为空。"],
            ["ChatResponse", "answer", "最终回答正文。", "字符串。"],
            ["ChatResponse", "intent", "识别出的意图。", "knowledge_query/data_analysis/fault_troubleshooting/error。"],
            ["ChatResponse", "sources", "知识、日志或案例来源。", "数组。"],
            ["ChatResponse", "sql", "Text2SQL 生成并执行的 SQL。", "仅数据分析场景通常有值。"],
            ["ChatResponse", "diagnostics", "故障诊断元数据。", "包含脚本输出、案例命中、附件等。"],
        ],
        [1.2, 1.35, 3.0, 1.0],
    )
    doc.add_heading("14.1 SSE 事件设计", 2)
    add_table(
        doc,
        ["事件", "触发时机", "data 结构", "前端处理"],
        [
            ["intent", "意图识别完成后立即发送。", "{type: intent}", "展示当前处理类型或记录消息元数据。"],
            ["token", "回答文本分块输出。", "{text: chunk}", "追加到助手消息内容。"],
            ["done", "回答完成后发送。", "{intent, sources, sql, diagnostics}", "补充来源、SQL 和诊断详情。"],
            ["error", "流式处理异常时发送。", "{message/error}", "展示错误提示并结束加载状态。"],
        ],
        [0.9, 1.9, 2.1, 2.0],
    )

    doc.add_heading("15 数据库与文件存储详细设计", 1)
    doc.add_heading("15.1 配置数据库", 2)
    add_table(
        doc,
        ["表名", "字段", "说明"],
        [
            ["datasource_configs", "id、name、type、config_json、is_active、created_at、updated_at", "保存数据源配置，敏感连接字段位于 JSON 中并加密。"],
            ["llm_provider_configs", "id、name、provider_type、api_key_encrypted、base_url、model、temperature、max_tokens、is_primary、created_at、updated_at", "保存模型提供商配置。"],
        ],
        [1.35, 3.65, 1.5],
    )
    doc.add_heading("15.2 故障案例数据库", 2)
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ["case_id", "TEXT PRIMARY KEY", "case_<uuid> 形式的稳定编号。"],
            ["query", "TEXT", "原始用户问题。"],
            ["answer", "TEXT", "故障排查生成的回答。"],
            ["symptoms", "TEXT(JSON)", "故障症状数组。"],
            ["root_cause", "TEXT", "根因摘要。"],
            ["solution", "TEXT", "处理方案摘要。"],
            ["evidence", "TEXT(JSON)", "日志、知识或脚本证据。"],
            ["status", "TEXT", "pending、resolved、auto_saved 等状态。"],
            ["category", "TEXT", "案例分类。"],
            ["tokens", "TEXT(JSON)", "用于相似匹配的 token 集合。"],
            ["created_at/updated_at", "TEXT", "UTC 时间戳。"],
        ],
        [1.55, 1.55, 3.4],
    )
    doc.add_heading("15.3 日志 Metadata", 2)
    add_table(
        doc,
        ["字段", "来源", "说明"],
        [
            ["file_id", "上传或本地发现", "上传日志为 log_<uuid>，本地日志为 local_<hash>。"],
            ["filename", "文件名", "展示用安全文件名。"],
            ["stored_path", "保存路径", "预览时必须校验路径在允许目录中。"],
            ["source", "uploaded/runtime/seed", "区分用户上传日志、运行日志和样例日志。"],
            ["category", "自动或人工", "用于日志管理工作台筛选。"],
            ["severity", "分析结果", "根据错误、警告和关键模式推断。"],
            ["analysis", "analyze_text", "包含 error_count、warning_count、patterns、summary。"],
        ],
        [1.45, 1.65, 3.4],
    )

    doc.add_heading("16 关键算法与规则", 1)
    doc.add_heading("16.1 意图分类规则", 2)
    add_table(
        doc,
        ["优先级", "规则", "结果"],
        [
            ["1", "存在 .log 或“日志文件/分析日志”等模式。", "fault_troubleshooting，置信度 0.95。"],
            ["2", "存在 select/from/limit、下划线表名或查询数据模式。", "data_analysis，置信度 0.95。"],
            ["3", "匹配故障、异常、CPU/内存/磁盘高、超时、OOM 等。", "累计故障排查分数。"],
            ["4", "匹配如何、命令、步骤、查看、检查等。", "累计知识查询分数。"],
            ["5", "无规则命中。", "默认 knowledge_query，置信度 0.5。"],
            ["6", "规则置信度不足 0.8。", "调用 LLM 分类；失败时使用规则结果。"],
        ],
        [0.8, 4.25, 1.45],
    )
    doc.add_heading("16.2 SQL 校验规则", 2)
    add_table(
        doc,
        ["校验项", "规则", "失败处理"],
        [
            ["危险关键字", "DROP、DELETE、INSERT、UPDATE、ALTER、TRUNCATE、CREATE、EXEC、GRANT、REVOKE、UNION 等禁止。", "抛出 DangerousSQLError。"],
            ["语句类型", "必须以 SELECT 开头。", "抛出只允许 SELECT 查询错误。"],
            ["LIMIT", "必须包含 LIMIT。", "抛出 SELECT 必须包含 LIMIT。"],
            ["LIMIT 上限", "LIMIT 数值不得超过 200。", "抛出 LIMIT 值过大。"],
            ["注入模式", "检测 OR 1=1、注释截断和拼接 SELECT 等。", "抛出可能 SQL 注入错误。"],
        ],
        [1.3, 3.55, 1.65],
    )
    doc.add_heading("16.3 相似案例匹配", 2)
    add_numbered(
        doc,
        [
            "保存案例时从 query、symptoms、root_cause、solution 中提取英文、数字和中文词片段。",
            "查询相似案例时同样提取 query 和 symptoms token。",
            "对最近 200 条 resolved/auto_saved 案例计算 token overlap 分数。",
            "如果症状 token 存在，则用症状重合度提高匹配分数。",
            "分数大于 min_score（默认 0.88）时返回最佳案例。",
        ],
    )

    doc.add_heading("17 前端状态与交互详细设计", 1)
    add_table(
        doc,
        ["区域", "状态", "交互细节"],
        [
            ["聊天输入区", "输入内容、上传附件、发送中状态。", "发送时禁用重复提交，附件上传后把 file_id 写入 attachments。"],
            ["消息列表", "用户消息、助手消息、流式累积内容。", "token 事件追加文本，done 事件补齐 sources/sql/diagnostics。"],
            ["知识库页面", "目录树、文件列表、预览详情。", "按目录过滤文件，上传后刷新树和列表。"],
            ["日志案例页面", "分类树、筛选条件、主列表、详情抽屉。", "日志和案例使用同类管理工作台结构，支持分类聚合。"],
            ["配置页面", "表单、连接测试、主力标记。", "保存前校验必填项，测试失败展示错误。"],
            ["索引页面", "collection 状态、重建按钮、加载状态。", "重建完成后刷新状态。"],
        ],
        [1.35, 1.75, 3.4],
    )

    doc.add_heading("18 测试用例设计", 1)
    add_table(
        doc,
        ["测试编号", "测试目标", "输入/操作", "预期结果"],
        [
            ["T-01", "日志文件名意图识别", "输入“分析 error.log”。", "IntentType 为 fault_troubleshooting。"],
            ["T-02", "下划线表名不误判日志", "输入数据表查询问题。", "IntentType 为 data_analysis。"],
            ["T-03", "日志附件契约", "attachments 包含 type=log。", "Orchestrator 强制故障排查。"],
            ["T-04", "SQL 写操作拦截", "DELETE FROM users LIMIT 1。", "抛出 DangerousSQLError。"],
            ["T-05", "SQL LIMIT 上限", "SELECT * FROM t LIMIT 1000。", "校验失败。"],
            ["T-06", "日志脱敏", "日志包含 password=abc。", "预览内容不出现 abc 明文。"],
            ["T-07", "本地日志发现", "source_dirs 下存在 syslog_sample.log。", "list_logs 能返回该文件。"],
            ["T-08", "案例相似匹配", "保存 resolved 案例后输入相似症状。", "find_similar 返回案例。"],
            ["T-09", "脚本名校验", "上传 run.sh。", "拒绝，必须 check_*.sh/py。"],
            ["T-10", "知识库文件限制", "上传不支持扩展名。", "接口拒绝。"],
            ["T-11", "索引重建", "调用知识库重建。", "返回重建成功或明确错误。"],
            ["T-12", "前端类型检查", "运行 vue-tsc。", "无 TypeScript 类型错误。"],
        ],
        [0.85, 1.55, 2.25, 2.4],
    )

    doc.add_heading("19 详细设计可追溯矩阵", 1)
    add_table(
        doc,
        ["需求编号", "设计落点", "测试落点"],
        [
            ["FR-01/02", "chat.py、Orchestrator.process/process_stream、前端 ChatView。", "test_orchestrator_contract.py、手工 SSE 演示。"],
            ["FR-03/04", "LogUploadService.resolve_mentioned_logs、Orchestrator._resolve_intent_for_attachments。", "test_chat_attachments_contract.py、test_intent.py。"],
            ["FR-05/06/07", "KnowledgeService、Retriever、IndexService。", "test_management_services.py、知识库页面演示。"],
            ["FR-08/09/10/11", "config routes、DataSourceFactory、Text2SQLGenerator、SQLValidator。", "test_text2sql.py、test_datasource_upload.py。"],
            ["FR-12/13/14", "LogUploadService、uploads/incidents routes、LogsCasesView。", "test_log_upload_service.py、日志页面演示。"],
            ["FR-15/16/17", "IncidentCaseMemory、incidents routes。", "test_incident_case_memory.py。"],
            ["FR-18/19", "DiagnosticService、diagnostics routes、ScriptExecutor。", "test_management_services.py、诊断页面演示。"],
            ["FR-20", "IndexService、indexes routes、IndexManagementView。", "test_management_services.py、索引页面演示。"],
        ],
        [1.0, 3.4, 2.1],
    )

def build_requirements_doc():
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "OpsAgent 中期验收 - 需求分析")
    add_cover(doc, "OpsAgent 智能运维助手", "需求分析说明书", "需求分析文档")
    add_toc_placeholder(doc, [
        "1 项目背景与建设目标",
        "2 用户角色与业务场景",
        "3 总体业务流程",
        "4 功能性需求",
        "5 非功能性需求",
        "6 中期功能演示计划",
        "7 验收标准与后续完善",
        "8 详细业务需求分析",
        "9 用例规约",
        "10 非功能需求量化指标",
        "11 中期验收材料清单",
    ])

    doc.add_heading("1 项目背景与建设目标", 1)
    doc.add_paragraph(
        "随着企业 IT 系统规模扩大，运维人员需要同时处理知识查询、数据统计、日志分析、故障排查和故障复盘等任务。"
        "传统运维工具往往把知识库、数据库查询、日志平台和诊断脚本分散在不同系统中，导致问题定位链路长、历史经验复用不足。"
        "OpsAgent 旨在构建基于大语言模型的智能运维助手，通过 Web 聊天界面把 RAG 知识检索、Text2SQL 数据分析、日志故障排查、诊断脚本和案例沉淀整合到统一工作流中。"
    )
    doc.add_heading("1.1 开题意见完善情况", 2)
    add_table(
        doc,
        ["开题阶段关注点", "完善措施", "中期体现"],
        [
            ["需求边界需更清晰", "将系统拆分为智能对话、知识库、数据源、日志案例、诊断工具、索引管理和模型配置七个业务模块。", "需求文档按模块列出功能、输入输出、优先级和验收标准。"],
            ["安全约束需明确", "补充 Text2SQL 只读限制、脚本白名单、路径穿越防护、密钥加密和日志脱敏。", "概要和详细设计文档均单列安全设计章节。"],
            ["演示功能需可闭环", "中期演示选择知识问答、Text2SQL、日志上传分析、案例复用和管理页面。", "需求文档列出功能演示清单和当前完成状态。"],
            ["文档需按验收格式准备", "采用 Word 文档、封面、目录、章节编号、表格、图示和版本信息。", "本次交付三份 .docx 文档。"],
        ],
        [1.55, 2.7, 2.15],
    )

    doc.add_heading("2 用户角色与业务场景", 1)
    add_table(
        doc,
        ["角色", "主要目标", "典型操作"],
        [
            ["一线运维人员", "快速获取处理步骤、定位日志异常并执行安全诊断。", "上传日志、询问故障原因、查看推荐处理方案、运行白名单脚本。"],
            ["运维负责人", "管理知识库、故障案例和索引状态，推动经验复用。", "维护知识文件、分类案例、重建索引、查看历史案例。"],
            ["数据分析人员", "通过自然语言查询运维数据，降低 SQL 门槛。", "配置数据源、提出统计问题、查看生成 SQL 与结果摘要。"],
            ["系统管理员", "配置模型提供商、数据源和安全策略。", "管理 LLM Provider、测试连接、设置主力模型。"],
        ],
        [1.3, 2.5, 2.7],
    )

    doc.add_heading("3 总体业务流程", 1)
    add_figure(doc, "flow.png", "图 3-1 OpsAgent 核心业务流程")
    add_numbered(
        doc,
        [
            "用户在 Web 聊天界面提交问题，必要时先上传日志附件。",
            "后端编排器执行意图分类，日志附件或 .log 文件名优先触发故障排查意图。",
            "系统按意图进入 RAG、Text2SQL 或故障排查处理链路。",
            "故障排查链路同时利用知识库、日志上下文、日志索引、诊断脚本和历史案例。",
            "LLM 对多源证据进行融合，返回结构化答案、来源、SQL 或诊断元数据。",
            "故障排查结果自动保存为案例，供后续相似问题复用。",
        ],
    )

    doc.add_heading("4 功能性需求", 1)
    add_table(
        doc,
        ["编号", "功能模块", "需求描述", "优先级", "中期状态"],
        [
            ["FR-01", "智能对话", "支持非流式和 SSE 流式聊天，返回 intent、token、done 等事件。", "高", "已实现雏形"],
            ["FR-02", "意图识别", "通过规则和 LLM 两级识别知识查询、数据分析和故障排查。", "高", "已实现"],
            ["FR-03", "RAG 知识检索", "支持运维知识文档检索、来源展示和 LLM 综合回答。", "高", "已实现雏形"],
            ["FR-04", "Text2SQL", "支持自然语言生成 SQL，执行只读查询并总结结果。", "高", "已实现雏形"],
            ["FR-05", "日志上传与分析", "支持日志保存、脱敏预览、摘要分析、附件上下文进入诊断。", "高", "已实现"],
            ["FR-06", "故障案例复用", "故障排查答案自动保存为案例，相似问题优先匹配历史方案。", "中", "已实现雏形"],
            ["FR-07", "知识库管理", "支持 .md/.txt 文件上传、目录分类、预览、删除和重建索引。", "中", "已实现"],
            ["FR-08", "诊断工具", "仅允许执行 approved 目录脚本，支持上传待启用、预览、启用、运行。", "中", "已实现雏形"],
            ["FR-09", "索引管理", "展示知识库、日志、案例 collection 状态，支持重建和清理。", "中", "已实现雏形"],
            ["FR-10", "配置管理", "支持数据源、LLM Provider 的增删改查、连接测试和主力配置。", "高", "已实现"],
        ],
        [0.55, 1.05, 3.25, 0.75, 0.9],
    )

    doc.add_heading("5 非功能性需求", 1)
    add_bullets(
        doc,
        [
            "安全性：SQL 仅允许 SELECT，禁止写操作；诊断脚本仅能从 scripts/approved 执行；路径解析限制在受控目录内。",
            "可用性：前端采用 Element Plus 构建中文管理界面，核心功能支持表单校验、状态反馈和错误提示。",
            "可维护性：后端按 api、core、models、data、utils 分层，前端按 api、components、stores、views、types 组织。",
            "可扩展性：LLM Provider、数据源类型、索引 collection 和诊断脚本均采用可扩展的服务层设计。",
            "可测试性：pytest 覆盖意图分类、Text2SQL 安全、日志上传、管理服务、案例记忆和编排契约。",
        ],
    )

    doc.add_heading("6 中期功能演示计划", 1)
    add_table(
        doc,
        ["演示项", "演示输入/操作", "预期结果", "完成度"],
        [
            ["知识问答", "询问 Linux 磁盘、Nginx、MySQL 处理步骤。", "返回步骤化答案并展示知识来源。", "可演示"],
            ["自然语言查数", "询问告警、服务器或工单统计。", "生成安全 SELECT SQL，执行并总结。", "可演示"],
            ["日志故障排查", "上传日志或指定 ops_agent_YYYY-MM-DD.log。", "自动进入故障排查，输出错误摘要和建议。", "可演示"],
            ["日志与案例管理", "筛选日志、预览脱敏内容、分类案例。", "展示工作台式列表与详情。", "可演示"],
            ["诊断脚本", "运行磁盘、CPU、内存、服务检查脚本。", "只执行白名单脚本，输出受限长度结果。", "可演示"],
            ["索引管理", "重建知识库、日志和案例索引。", "展示 collection 状态和重建结果。", "可演示"],
        ],
        [1.1, 2.1, 2.2, 0.9],
    )

    doc.add_heading("7 验收标准与后续完善", 1)
    add_table(
        doc,
        ["类别", "中期验收标准", "最终验收待完善"],
        [
            ["文档", "三份 Word 文档具备规范目录、章节、表格和图示。", "补充测试截图、部署截图、最终用户手册。"],
            ["功能", "核心页面可访问，聊天、配置、日志、知识库等具备演示链路。", "增强真实数据源联调、异常场景覆盖和体验细节。"],
            ["质量", "关键服务具备 pytest 测试，SQL 和路径安全有基础约束。", "补充端到端测试、性能测试和最终部署说明。"],
        ],
        [1.2, 2.85, 2.75],
    )
    add_requirements_expansion(doc)
    doc.save(OUT / "OpsAgent_中期验收_需求分析.docx")


def build_outline_doc():
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "OpsAgent 中期验收 - 概要设计")
    add_cover(doc, "OpsAgent 智能运维助手", "概要设计说明书", "概要设计文档")
    add_toc_placeholder(doc, [
        "1 设计目标与原则",
        "2 系统总体架构",
        "3 功能模块设计",
        "4 技术架构与数据流",
        "5 部署与运行设计",
        "6 接口概要设计",
        "7 安全概要设计",
        "8 测试与质量保障",
        "9 架构设计决策",
        "10 模块协作关系",
        "11 运行时数据与目录规划",
        "12 接口矩阵",
        "13 关键流程概要",
        "14 概要层风险与对策",
    ])

    doc.add_heading("1 设计目标与原则", 1)
    add_bullets(
        doc,
        [
            "分层清晰：API 层、核心编排层、模型能力层、数据层和前端展示层职责分离。",
            "闭环运维：从问题输入、证据检索、结果融合到案例沉淀形成完整工作流。",
            "安全优先：Text2SQL、脚本执行、文件路径和敏感信息展示均设置显式边界。",
            "可演示可扩展：中期优先完成稳定雏形，最终阶段继续补充精度、体验和部署细节。",
        ],
    )

    doc.add_heading("2 系统总体架构", 1)
    add_figure(doc, "architecture.png", "图 2-1 OpsAgent 系统总体架构图")
    doc.add_paragraph(
        "系统采用前后端分离与后端统一托管构建产物的组合方式。开发阶段由 Vite 提供前端开发服务器并代理 /api 到 FastAPI；"
        "构建后由 FastAPI 挂载 Vue SPA 的静态资源。后端服务负责聊天编排、配置管理、知识库管理、日志案例管理、诊断脚本和索引管理。"
    )

    doc.add_heading("3 功能模块设计", 1)
    add_figure(doc, "modules.png", "图 3-1 功能模块划分图")
    add_table(
        doc,
        ["模块", "前端页面", "后端服务/接口", "概要说明"],
        [
            ["智能对话", "ChatView", "/api/chat、/api/chat/stream、Orchestrator", "承接自然语言输入并根据意图分发到不同能力链路。"],
            ["知识库管理", "KnowledgeView", "/api/knowledge/*、KnowledgeService", "管理知识文档、目录分类、预览和索引重建。"],
            ["日志与案例", "LogsCasesView", "/api/uploads/*、/api/incidents/*", "统一日志目录、脱敏预览、案例分类和历史复用。"],
            ["诊断工具", "DiagnosticsView", "/api/diagnostics/*、DiagnosticService", "管理白名单脚本和待启用脚本，执行时限制超时和输出长度。"],
            ["索引管理", "IndexManagementView", "/api/indexes/*、IndexService", "查看 Milvus collection 状态，重建知识、日志和案例索引。"],
            ["数据源配置", "DataSourceView", "/api/config/datasources/*", "管理 MySQL、ClickHouse、Excel/CSV 数据源。"],
            ["模型配置", "LLMConfigView", "/api/config/llm/*", "管理 OpenAI 兼容和 DashScope 提供商。"],
        ],
        [1.05, 1.25, 2.1, 2.4],
    )

    doc.add_heading("4 技术架构与数据流", 1)
    add_table(
        doc,
        ["层次", "技术/组件", "职责"],
        [
            ["前端层", "Vue 3、TypeScript、Vite、Element Plus、Pinia、Vue Router", "页面展示、状态管理、API 调用和 SSE 流式消息呈现。"],
            ["API 层", "FastAPI、Pydantic、Uvicorn", "路由注册、请求校验、静态资源托管、CORS/API Key 中间件。"],
            ["核心层", "IntentClassifier、TaskRouter、ResponseFusion、Orchestrator", "意图识别、任务分发、多源证据融合和流式输出。"],
            ["模型能力层", "LLM Client、BGE Embedder、RAG、Text2SQL、脚本执行器", "模型调用、向量检索、SQL 生成校验、诊断工具调用。"],
            ["数据层", "SQLite、Milvus Lite、文件目录、MySQL/ClickHouse/CSV", "配置存储、向量索引、知识/日志/案例文件和业务数据查询。"],
        ],
        [1.0, 2.35, 3.05],
    )

    doc.add_heading("5 部署与运行设计", 1)
    add_figure(doc, "deployment.png", "图 5-1 部署与运行视图")
    add_bullets(
        doc,
        [
            "后端启动命令：uvicorn ops_agent.api.main:app --reload --port 8080。",
            "前端开发命令：cd frontend && npm run dev，Vite 将 /api 代理到 localhost:8080。",
            "构建输出：cd frontend && npm run build，产物输出到 ops_agent/api/static/dist/。",
            "运行数据：data/uploads、data/vectors、data/app_config.db、data/incident_cases.db 为本地运行产物，不纳入源码提交。",
        ],
    )

    doc.add_heading("6 接口概要设计", 1)
    add_table(
        doc,
        ["接口分组", "主要路径", "说明"],
        [
            ["聊天", "/api/chat、/api/chat/stream", "非流式和 SSE 流式对话。"],
            ["配置", "/api/config/datasources/*、/api/config/llm/*", "数据源与大模型 Provider 管理。"],
            ["知识库", "/api/knowledge/files、/tree、/folders、/upload、/reindex", "文件型知识库管理。"],
            ["日志上传", "/api/uploads/logs", "聊天日志附件上传。"],
            ["日志案例", "/api/incidents/logs、/api/incidents/incidents", "日志目录、预览、分类、案例状态与分类。"],
            ["诊断", "/api/diagnostics/scripts、/pending、/upload、/run", "脚本管理与安全执行。"],
            ["索引", "/api/indexes/status、/knowledge/rebuild、/logs/rebuild、/cases/rebuild", "索引状态与重建。"],
        ],
        [1.15, 2.35, 3.0],
    )

    doc.add_heading("7 安全概要设计", 1)
    add_bullets(
        doc,
        [
            "SQL 安全：禁止 DROP、DELETE、INSERT、UPDATE、ALTER、TRUNCATE、CREATE 等写操作，要求 SELECT 且 LIMIT 不超过 200。",
            "脚本安全：仅执行 scripts/approved 中的脚本，上传脚本进入待启用区，执行超时 30 秒，输出截断 5000 字符。",
            "路径安全：知识库、日志和脚本接口必须将解析路径限制在对应根目录内。",
            "密钥安全：SQLite 配置中的 API Key 和数据库密码使用 Fernet 加密保存。",
            "日志安全：日志预览和上下文摘要对 password、token、secret、api_key、Authorization Bearer 等字段脱敏。",
        ],
    )

    doc.add_heading("8 测试与质量保障", 1)
    add_table(
        doc,
        ["测试文件", "覆盖重点"],
        [
            ["test_intent.py", "意图分类、日志文件名优先故障排查、实体提取。"],
            ["test_text2sql.py", "SQL 安全校验、LIMIT 约束和危险关键字拦截。"],
            ["test_log_upload_service.py", "日志保存、目录发现、脱敏、非法扩展名和附件解析。"],
            ["test_incident_case_memory.py", "故障案例保存、分类、相似匹配与状态管理。"],
            ["test_management_services.py", "知识库、索引、诊断脚本等管理服务契约。"],
            ["test_orchestrator_contract.py", "编排器返回结构、附件强制故障排查和 SSE 元数据。"],
        ],
        [1.75, 4.75],
    )
    add_outline_expansion(doc)
    doc.save(OUT / "OpsAgent_中期验收_概要设计.docx")


def build_detail_doc():
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "OpsAgent 中期验收 - 详细设计")
    add_cover(doc, "OpsAgent 智能运维助手", "详细设计说明书", "详细设计文档")
    add_toc_placeholder(doc, [
        "1 包结构与职责设计",
        "2 核心数据对象设计",
        "3 聊天编排详细设计",
        "4 意图分类详细设计",
        "5 RAG 知识检索详细设计",
        "6 Text2SQL 详细设计",
        "7 日志与故障排查详细设计",
        "8 诊断脚本详细设计",
        "9 前端页面详细设计",
        "10 API 详细设计摘录",
        "11 安全边界详细设计",
        "12 中期未完成项与最终阶段计划",
        "13 核心类与方法设计",
        "14 请求与响应结构设计",
        "15 数据库与文件存储详细设计",
        "16 关键算法与规则",
        "17 前端状态与交互详细设计",
        "18 测试用例设计",
        "19 详细设计可追溯矩阵",
    ])

    doc.add_heading("1 包结构与职责设计", 1)
    add_table(
        doc,
        ["目录/文件", "职责"],
        [
            ["ops_agent/api/main.py", "创建 FastAPI 应用，初始化配置数据库，注册路由，挂载前端静态资源。"],
            ["ops_agent/api/routes/*.py", "提供聊天、配置、知识库、日志、案例、诊断和索引管理接口。"],
            ["ops_agent/core/orchestrator.py", "系统核心编排器，负责意图识别、任务分发、故障排查和流式输出。"],
            ["ops_agent/core/intent/classifier.py", "规则快速通道与 LLM 精准分类，抽取 IP、主机、服务、端口等实体。"],
            ["ops_agent/models/rag/*", "知识库检索、日志索引和向量查询。"],
            ["ops_agent/models/text2sql/*", "Schema 管理、SQL 生成和 SQL 安全校验。"],
            ["ops_agent/models/uploads/log_upload_service.py", "上传日志、发现本地日志、脱敏预览、摘要分析和附件上下文生成。"],
            ["ops_agent/models/troubleshooting/case_memory.py", "故障案例存储、分类、状态更新和相似案例匹配。"],
            ["frontend/src/views/*.vue", "七个主要页面：聊天、数据源、知识库、日志案例、诊断、索引、模型配置。"],
        ],
        [2.25, 4.25],
    )

    doc.add_heading("2 核心数据对象设计", 1)
    add_figure(doc, "data_model.png", "图 2-1 核心数据对象关系")
    add_table(
        doc,
        ["对象", "关键字段", "说明"],
        [
            ["LLMProvider", "id、name、provider_type、api_key、base_url、model、is_primary", "管理 DeepSeek、DashScope 或其他 OpenAI 兼容模型。"],
            ["DataSource", "id、type、host/path、database、credentials、is_active", "管理 MySQL、ClickHouse、Excel/CSV 数据源。"],
            ["LogMetadata", "file_id、filename、source、category、severity、tags、analysis、stored_path", "描述上传日志和本地发现日志。"],
            ["IncidentCase", "case_id、symptoms、root_cause、solution、evidence、status、category", "保存故障排查结果并支持复用。"],
            ["KnowledgeFile", "file_id、relative_path、folder、mtime、index_status", "描述知识库文件与索引状态。"],
        ],
        [1.25, 2.65, 2.6],
    )

    doc.add_heading("3 聊天编排详细设计", 1)
    doc.add_paragraph("Orchestrator 同时提供 process 和 process_stream 两种入口。process_stream 输出 intent、token、done 三类主要 SSE 事件。")
    add_numbered(
        doc,
        [
            "合并显式上传附件和用户文本中提到的日志文件名。",
            "调用 IntentClassifier 执行规则与 LLM 意图识别。",
            "如果附件 type 为 log，则强制将意图切换为 fault_troubleshooting。",
            "TaskRouter 根据意图调用知识问答、数据分析或故障排查处理器。",
            "处理器返回 answer、sources、sql、diagnostics 等字段。",
            "流式接口按固定 chunk 输出 token，并在 done 事件返回元数据。",
        ],
    )

    doc.add_heading("4 意图分类详细设计", 1)
    add_table(
        doc,
        ["规则类别", "识别内容", "处理策略"],
        [
            ["日志分析提示", ".log、日志文件、分析日志、error.log、access.log", "置信度 0.95，直接判定为 fault_troubleshooting。"],
            ["数据查询提示", "select/from/limit、下划线表名、查询/统计/显示数据等", "置信度 0.95，判定为 data_analysis。"],
            ["故障排查词", "故障、异常、报错、CPU 高、内存满、超时、OOM 等", "累计规则分数后判定故障排查。"],
            ["知识查询词", "如何、怎么、命令、步骤、查看、检查等", "默认安全兜底为 knowledge_query。"],
            ["LLM 分类", "规则置信度不足时调用模型输出 JSON", "解析 intent、confidence、entities，失败时回退规则结果。"],
        ],
        [1.15, 2.75, 2.6],
    )

    doc.add_heading("5 RAG 知识检索详细设计", 1)
    add_bullets(
        doc,
        [
            "知识文件来源于 data/knowledge，管理接口只允许 .md 和 .txt。",
            "DocumentLoader 负责文档加载和分块，Embedder 生成 BGE 向量。",
            "VectorStore 使用 Milvus Lite 保存 collection，检索结果返回标题、来源文件和片段。",
            "ResponseFusion 将用户问题和检索上下文交给 LLM，生成面向运维人员的步骤化回答。",
            "IndexService 维护知识库索引状态，文件修改时间晚于索引时间时标记为待重建。",
        ],
    )

    doc.add_heading("6 Text2SQL 详细设计", 1)
    add_table(
        doc,
        ["步骤", "设计说明"],
        [
            ["数据源选择", "优先使用活跃数据源，用户传入 datasource_id 时可临时覆盖。"],
            ["Schema 管理", "SchemaManager 获取表结构并缓存，历史对话可补充上一次提到的表名。"],
            ["SQL 生成", "Text2SQLGenerator 将自然语言、schema 和约束 prompt 发送给 LLM。"],
            ["安全校验", "SQLValidator 禁止写操作、要求 SELECT、要求 LIMIT，LIMIT 最大 200。"],
            ["执行与总结", "数据源执行查询后，ResponseFusion 根据结果生成自然语言总结。"],
        ],
        [1.25, 5.25],
    )

    doc.add_heading("7 日志与故障排查详细设计", 1)
    add_table(
        doc,
        ["子功能", "详细设计"],
        [
            ["日志保存", "上传文件校验扩展名和大小，保存到 data/uploads/logs/files，并写入 metadata JSON。"],
            ["本地日志发现", "从运行日志和样例日志目录发现 .log/.txt/.out/.gz 文件，生成稳定 local_<hash> file_id。"],
            ["日志脱敏", "预览和上下文摘要对 password、token、secret、api_key、Authorization Bearer 进行替换。"],
            ["附件上下文", "将错误数、警告数、关键模式和关键片段渲染为 LLM 诊断上下文。"],
            ["案例复用", "故障排查前先查询相似案例，命中后返回历史根因和解决方案。"],
            ["案例沉淀", "未命中案例时组合知识库、日志、脚本结果生成答案，并自动保存为 IncidentCase。"],
        ],
        [1.35, 5.15],
    )

    doc.add_heading("8 诊断脚本详细设计", 1)
    add_bullets(
        doc,
        [
            "脚本来源分为 approved 与 pending，用户上传脚本先进入 pending。",
            "启用脚本时校验文件名、扩展名和路径，并移动到 scripts/approved。",
            "执行器只接收 approved 目录中的脚本名，不接受任意命令行输入。",
            "脚本执行设置 30 秒超时，输出最多保留 5000 字符，避免长输出影响聊天上下文。",
            "故障排查链路按用户问题中的 CPU、内存、磁盘、服务等实体选择相关脚本。",
        ],
    )

    doc.add_heading("9 前端页面详细设计", 1)
    add_table(
        doc,
        ["页面", "路由", "主要交互"],
        [
            ["ChatView", "/", "发送问题、上传日志附件、展示流式回答、来源和 SQL。"],
            ["DataSourceView", "/datasources", "创建、编辑、测试、激活数据源。"],
            ["KnowledgeView", "/knowledge", "目录树、上传文件、预览文件、重建索引。"],
            ["LogsCasesView", "/logs-cases", "日志和案例分类、筛选、预览、状态更新。"],
            ["DiagnosticsView", "/diagnostics", "查看脚本、上传待启用脚本、启用、运行、删除。"],
            ["IndexManagementView", "/indexes", "查看 collection 状态，重建知识/日志/案例索引。"],
            ["LLMConfigView", "/llm", "管理模型提供商，测试连接，设置主力模型。"],
        ],
        [1.45, 1.1, 4.0],
    )

    doc.add_heading("10 API 详细设计摘录", 1)
    add_table(
        doc,
        ["方法", "路径", "功能"],
        [
            ["POST", "/api/chat", "提交非流式聊天请求。"],
            ["POST", "/api/chat/stream", "提交 SSE 流式聊天请求。"],
            ["GET/POST/PUT/DELETE", "/api/config/datasources", "数据源 CRUD、测试、激活和表结构查询。"],
            ["GET/POST/PUT/DELETE", "/api/config/llm", "模型 Provider CRUD、测试和设为主力。"],
            ["GET/POST/DELETE", "/api/knowledge/*", "知识文件、目录、上传、预览、删除和重建索引。"],
            ["POST", "/api/uploads/logs", "上传聊天日志附件。"],
            ["GET/PUT/DELETE", "/api/incidents/logs/*", "日志列表、预览、分类和删除。"],
            ["GET/PUT/DELETE", "/api/incidents/incidents/*", "故障案例列表、详情、状态、分类和删除。"],
            ["GET/POST/DELETE", "/api/diagnostics/*", "脚本列表、上传、启用、运行和删除。"],
            ["GET/POST", "/api/indexes/*", "索引状态、重建和清理。"],
        ],
        [1.0, 2.3, 3.2],
    )

    doc.add_heading("11 安全边界详细设计", 1)
    add_table(
        doc,
        ["风险", "防护措施", "落点"],
        [
            ["SQL 注入或误写库", "危险关键字拦截、只允许 SELECT、强制 LIMIT ≤ 200。", "SQLValidator"],
            ["任意命令执行", "仅按脚本名执行 approved 目录脚本，上传脚本需启用。", "ScriptExecutor、DiagnosticService"],
            ["路径穿越", "解析后的路径必须位于知识库、日志或脚本根目录内。", "KnowledgeService、LogUploadService、DiagnosticService"],
            ["密钥泄露", "配置数据库加密保存密钥，日志预览脱敏。", "ConfigService、LogUploadService"],
            ["大模型不可用", "模型客户端支持 DeepSeek 与 DashScope 自动回退。", "UnifiedLLMClient"],
        ],
        [1.25, 3.25, 2.0],
    )

    doc.add_heading("12 中期未完成项与最终阶段计划", 1)
    add_table(
        doc,
        ["事项", "当前状态", "后续计划"],
        [
            ["真实生产数据联调", "已有多数据源框架和测试接口。", "补充 MySQL/ClickHouse 真实样例库联调和截图。"],
            ["知识库内容规模", "已有 Linux、MySQL、监控、服务运维知识样例。", "扩展更多场景文档并优化检索召回质量。"],
            ["端到端验证", "已有 pytest 服务测试和前端类型检查命令。", "补充浏览器端演示脚本和最终验收测试记录。"],
            ["部署说明", "已有开发和构建流程。", "补充一键启动脚本、环境变量说明和部署截图。"],
            ["用户手册", "本次主要交付分析和设计文档。", "最终阶段增加用户操作手册和管理员手册。"],
        ],
        [1.35, 2.25, 2.9],
    )
    add_detail_expansion(doc)
    doc.save(OUT / "OpsAgent_中期验收_详细设计.docx")


def build_all():
    OUT.mkdir(parents=True, exist_ok=True)
    build_figures()
    build_requirements_doc()
    build_outline_doc()
    build_detail_doc()


if __name__ == "__main__":
    build_all()
    for path in sorted(OUT.glob("*.docx")):
        print(path)
