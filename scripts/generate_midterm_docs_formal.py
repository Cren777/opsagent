from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "build" / "midterm_acceptance_formal"
FIG = OUT / "figures"

BLACK = RGBColor(0, 0, 0)
BORDER = "000000"


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        try:
            return ImageFont.truetype(item, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def set_east_asia(run, name: str = "宋体"):
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), BORDER)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def blacken_runs(doc: Document):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = BLACK
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = BLACK


def add_header_footer(doc: Document, label: str):
    section = doc.sections[0]
    p = section.header.paragraphs[0]
    p.text = label
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = BLACK
        set_east_asia(run, "宋体")

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("第 ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    fp.add_run(" 页")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = BLACK
                set_east_asia(r, "黑体")
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = text
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = BLACK
                    set_east_asia(r, "宋体")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_cover(doc: Document, title: str, doc_type: str):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = BLACK
    set_east_asia(r, "黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_type)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLACK
    set_east_asia(r, "黑体")

    doc.add_paragraph()
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["项目名称", "OpsAgent 智能运维助手"],
            ["文档类型", doc_type],
            ["阶段", "中期验收"],
            ["版本", "V1.0"],
            ["日期", "2026 年 6 月"],
        ],
        [1.7, 4.7],
    )
    doc.add_page_break()


def add_toc(doc: Document, entries: list[str]):
    doc.add_heading("目录", 1)
    for entry in entries:
        doc.add_paragraph(entry)
    doc.add_page_break()


def draw_er_diagram():
    FIG.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (2200, 1500), "white")
    d = ImageDraw.Draw(img)
    title_font = pil_font(42, True)
    box_font = pil_font(26, True)
    text_font = pil_font(21)
    d.text((60, 40), "OpsAgent 数据实体关系图（E-R 图）", fill="black", font=title_font)

    boxes = {
        "datasource": (90, 180, 560, 420, "datasource_configs", ["id PK", "name", "type", "config_json", "is_active"]),
        "llm": (855, 180, 1325, 420, "llm_provider_configs", ["id PK", "name", "provider_type", "api_key_encrypted", "is_primary"]),
        "knowledge": (1620, 180, 2090, 420, "knowledge_file", ["file_id PK", "relative_path", "filename", "size", "indexed"]),
        "log": (90, 680, 560, 940, "log_metadata", ["file_id PK", "filename", "source", "category", "analysis"]),
        "case": (855, 680, 1325, 940, "incident_cases", ["case_id PK", "query", "symptoms", "status", "category"]),
        "category": (1620, 680, 2090, 940, "category_registry", ["name PK", "pinned", "user_defined"]),
        "vector": (780, 1130, 1400, 1380, "Milvus collections", ["ops_knowledge", "ops_logs", "ops_incident_cases"]),
    }

    for x1, y1, x2, y2, name, fields in boxes.values():
        d.rectangle((x1, y1, x2, y2), outline="black", width=3)
        d.line((x1, y1 + 45, x2, y1 + 45), fill="black", width=2)
        d.text((x1 + 14, y1 + 10), name, fill="black", font=box_font)
        y = y1 + 62
        for field in fields:
            d.text((x1 + 18, y), field, fill="black", font=text_font)
            y += 32

    def center(key: str) -> tuple[int, int]:
        x1, y1, x2, y2, *_ = boxes[key]
        return ((x1 + x2) // 2, (y1 + y2) // 2)

    def point(key: str, side: str) -> tuple[int, int]:
        x1, y1, x2, y2, *_ = boxes[key]
        if side == "top":
            return ((x1 + x2) // 2, y1)
        if side == "bottom":
            return ((x1 + x2) // 2, y2)
        if side == "left":
            return (x1, (y1 + y2) // 2)
        return (x2, (y1 + y2) // 2)

    def line_between(start: tuple[int, int], end: tuple[int, int]):
        d.line((start, end), fill="black", width=2)

    def elbow(points: list[tuple[int, int]]):
        d.line(points, fill="black", width=2)

    # 关系线只表达实体关联；关系含义在文档正文表格中说明，避免图面拥挤。
    line_between(point("datasource", "right"), point("case", "left"))
    line_between(point("llm", "bottom"), point("case", "top"))
    line_between(point("log", "right"), point("case", "left"))
    line_between(point("category", "left"), point("case", "right"))
    elbow([point("category", "bottom"), (1855, 1020), (325, 1020), point("log", "bottom")])
    elbow([point("knowledge", "bottom"), (1855, 560), (1510, 560), (1510, 1080), (1400, 1080), point("vector", "top")])
    elbow([point("log", "bottom"), (325, 1060), (780, 1060), point("vector", "left")])
    line_between(point("case", "bottom"), point("vector", "top"))

    d.text((90, 1450), "说明：图中矩形表示数据实体或索引集合，连线表示配置、知识、日志、案例、分类与向量索引之间存在业务关联。", fill="black", font=pil_font(22))
    img.save(FIG / "er_diagram.png")


def add_er(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(FIG / "er_diagram.png"), width=Inches(6.6))
    c = doc.add_paragraph("图 4-1 OpsAgent 数据实体关系图")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_requirements_doc():
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "OpsAgent 中期验收 - 需求分析")
    add_cover(doc, "OpsAgent 智能运维助手", "需求分析说明书")
    add_toc(doc, [
        "1 引言",
        "2 项目背景与问题分析",
        "3 用户角色与业务场景",
        "4 功能需求",
        "5 用例规约",
        "6 非功能需求",
        "7 开题意见整改说明",
        "8 中期演示与验收范围",
    ])

    doc.add_heading("1 引言", 1)
    doc.add_paragraph("OpsAgent 是面向企业 IT 运维场景的智能运维助手。系统通过 Web 对话界面接收自然语言问题，结合知识库、业务数据源、日志文件、诊断脚本和历史案例，为运维人员提供知识问答、数据查询、日志分析、故障排查和经验复用能力。")
    doc.add_paragraph("本文档用于中期验收阶段说明系统需求范围、用户角色、功能需求、用例规约和验收依据。文档内容按照最终验收材料要求组织，后续阶段将在真实演示截图、部署记录和测试报告方面继续补充。")

    doc.add_heading("2 项目背景与问题分析", 1)
    add_table(
        doc,
        ["问题", "现状影响", "系统解决方式"],
        [
            ["运维知识分散", "故障处理步骤、命令和经验散落在文档或个人经验中，查询效率低。", "通过知识库管理与 RAG 检索统一提供步骤化问答。"],
            ["数据查询门槛高", "运维人员需要掌握 SQL 才能查询告警、服务器、工单等数据。", "通过 Text2SQL 把自然语言转换为只读 SQL，并返回结果摘要。"],
            ["日志分析链路长", "日志上传、脱敏、搜索、诊断往往需要多个工具协同。", "统一日志上传、目录发现、脱敏预览和故障上下文注入。"],
            ["历史案例复用不足", "相似故障重复排查，根因和方案难以沉淀。", "故障排查结果保存为案例，相似问题优先匹配历史处理方案。"],
            ["自动诊断存在风险", "脚本执行可能带来任意命令风险。", "只允许执行白名单目录中的诊断脚本，上传脚本需要启用后才能运行。"],
        ],
        [1.35, 2.55, 2.6],
    )

    doc.add_heading("3 用户角色与业务场景", 1)
    add_table(
        doc,
        ["角色", "职责", "主要使用场景"],
        [
            ["一线运维人员", "处理日常告警、日志异常和服务故障。", "智能问答、上传日志分析、运行诊断脚本、查看历史案例。"],
            ["运维负责人", "维护运维知识和故障复盘资料。", "管理知识库、日志分类、故障案例状态、索引重建。"],
            ["数据分析人员", "根据运维数据进行统计分析。", "配置数据源、使用自然语言查询数据、查看 SQL 和统计结果。"],
            ["系统管理员", "维护系统运行配置和外部服务连接。", "配置 LLM Provider、配置数据库连接、审核诊断脚本。"],
        ],
        [1.2, 2.2, 3.1],
    )

    doc.add_heading("4 功能需求", 1)
    add_table(
        doc,
        ["编号", "功能", "输入", "处理过程", "输出", "异常处理"],
        [
            ["FR-01", "智能对话", "query、history、datasource_id、attachments。", "后端编排器识别意图并分发到对应处理链路。", "answer、intent、sources、sql、diagnostics。", "处理异常时返回 intent=error 和错误说明。"],
            ["FR-02", "流式对话", "ChatRequest。", "通过 SSE 发送 intent、token、done 事件。", "前端逐步显示回答，并在完成后展示来源、SQL 或诊断信息。", "异常时返回 error 事件。"],
            ["FR-03", "知识库管理", "md/txt 文件、目录路径。", "校验文件类型和路径后保存，支持目录树、预览、删除和重建索引。", "文件列表、目录树、文件内容、索引结果。", "非法路径、非法扩展名、文件不存在返回错误。"],
            ["FR-04", "知识问答", "用户运维问题。", "检索知识库向量索引，组合上下文调用 LLM。", "步骤化答案和来源文件。", "无索引或无命中时返回可理解提示。"],
            ["FR-05", "数据源配置", "MySQL、ClickHouse、Excel/CSV 配置。", "保存配置，敏感字段加密，支持连接测试和激活。", "数据源配置列表、测试结果、表结构。", "连接失败返回错误原因。"],
            ["FR-06", "Text2SQL 数据分析", "自然语言查询和可选数据源。", "获取 schema、生成 SQL、执行安全校验、查询数据并总结。", "自然语言摘要、SQL、rows。", "无数据源、SQL 不安全或执行失败时返回错误。"],
            ["FR-07", "日志上传与预览", "日志文件名、分类、文件内容。", "校验扩展名和大小，保存文件与 metadata，提取错误摘要并脱敏。", "日志 metadata、分析摘要、脱敏预览内容。", "空文件、非法扩展名、超大小返回错误。"],
            ["FR-08", "故障排查", "query、日志附件、历史对话。", "合并日志上下文、知识库、日志索引、诊断脚本和历史案例。", "故障现象、可能原因、排查步骤、处理建议、diagnostics。", "案例或脚本不可用时跳过对应证据。"],
            ["FR-09", "故障案例管理", "case_id、status、category、query。", "保存、查询、筛选、更新状态和分类。", "案例列表、详情、分类统计。", "案例不存在返回 404。"],
            ["FR-10", "诊断脚本管理", "check_*.sh 或 check_*.py、脚本参数。", "脚本上传进入 pending，启用后进入 approved，执行时限制目录、超时和输出。", "脚本列表、预览、执行 stdout/stderr/exit_code。", "非法脚本名、脚本不存在或执行超时返回错误。"],
            ["FR-11", "索引管理", "collection 或重建类型。", "查看 Milvus collection 状态，重建知识、日志和案例索引。", "collection 名称、数量、状态、重建结果。", "collection 异常返回错误状态。"],
        ],
        [0.55, 1.05, 1.55, 2.15, 1.3, 1.25],
    )

    doc.add_heading("5 用例规约", 1)
    add_table(
        doc,
        ["用例", "参与者", "前置条件", "主流程", "成功结果"],
        [
            ["UC-01 智能知识问答", "一线运维人员", "知识库索引可用。", "输入运维问题，系统识别为知识查询并检索知识库。", "返回操作步骤和来源文件。"],
            ["UC-02 自然语言查数", "数据分析人员", "已配置并激活数据源。", "输入统计问题，系统生成 SQL，校验后执行。", "返回 SQL、查询结果和文字总结。"],
            ["UC-03 上传日志诊断", "一线运维人员", "日志文件合法。", "上传日志后发起诊断请求，系统强制进入故障排查。", "返回日志摘要、故障原因和处理建议。"],
            ["UC-04 管理知识库", "运维负责人", "进入知识库页面。", "创建目录、上传文件、预览文件、重建索引。", "知识文件可被问答检索。"],
            ["UC-05 复用故障案例", "一线运维人员", "案例库存在相似 resolved/auto_saved 案例。", "输入相似问题，系统先查询案例库。", "直接返回历史根因和解决方案。"],
            ["UC-06 执行诊断脚本", "系统管理员/运维人员", "脚本已在 approved 目录。", "选择脚本并输入参数执行。", "返回执行输出和退出码。"],
        ],
        [1.15, 1.15, 1.55, 2.35, 1.65],
    )

    doc.add_heading("6 非功能需求", 1)
    add_table(
        doc,
        ["类别", "需求说明", "实现依据"],
        [
            ["安全性", "Text2SQL 只允许 SELECT，必须包含 LIMIT，禁止危险关键字，LIMIT 最大 200。", "SQLValidator。"],
            ["安全性", "诊断脚本只能从 scripts/approved 执行，上传脚本必须通过文件名和路径校验。", "DiagnosticService、ScriptExecutor。"],
            ["安全性", "知识、日志、脚本路径均需限制在对应根目录内。", "KnowledgeService、LogUploadService、DiagnosticService。"],
            ["保密性", "API Key、数据库密码加密保存，日志预览脱敏 password、token、secret 等字段。", "config_service、LogUploadService。"],
            ["可维护性", "后端按 API、core、models、data 分层，前端按 api、views、components、types 分层。", "项目目录结构。"],
            ["可测试性", "关键能力具备 pytest 测试，覆盖意图、SQL、安全、日志、案例和管理服务。", "tests 目录。"],
        ],
        [1.1, 3.25, 2.15],
    )

    doc.add_heading("7 开题意见整改说明", 1)
    add_table(
        doc,
        ["开题意见", "整改内容"],
        [
            ["需求边界不够明确", "将功能拆分为对话、知识库、数据分析、日志案例、诊断工具、索引管理和配置管理，并明确输入输出。"],
            ["安全设计需要强化", "补充 SQL 只读校验、日志脱敏、路径限制、脚本白名单和配置加密。"],
            ["演示功能需要闭环", "设计了从日志上传、故障诊断、案例保存到相似案例复用的闭环演示。"],
            ["文档格式需符合验收要求", "以 Word 文档组织需求分析、概要设计和详细设计，包含接口表、字段表和 ER 图。"],
        ],
        [2.0, 4.5],
    )

    doc.add_heading("8 中期演示与验收范围", 1)
    add_table(
        doc,
        ["演示项", "操作", "验收点"],
        [
            ["智能对话", "在首页输入知识类问题。", "能返回答案、意图和来源。"],
            ["Text2SQL", "配置数据源后输入统计问题。", "能生成安全 SQL 并返回结果摘要。"],
            ["日志故障排查", "上传日志或指定日志文件名。", "能进入 fault_troubleshooting 并输出诊断信息。"],
            ["知识库管理", "上传 md/txt 文件并重建索引。", "文件可预览，索引状态更新。"],
            ["日志与案例管理", "筛选日志、查看脱敏预览、修改案例分类。", "列表和详情数据正确。"],
            ["诊断脚本管理", "上传待启用脚本、启用、执行。", "只允许 approved 脚本运行。"],
            ["索引管理", "重建知识、日志、案例索引。", "返回 collection 和 count。"],
        ],
        [1.4, 2.6, 2.5],
    )

    blacken_runs(doc)
    doc.save(OUT / "OpsAgent_中期验收_需求分析.docx")


def build_outline_doc():
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "OpsAgent 中期验收 - 概要设计")
    add_cover(doc, "OpsAgent 智能运维助手", "概要设计说明书")
    add_toc(doc, [
        "1 设计目标",
        "2 系统总体结构",
        "3 后端分层设计",
        "4 前端结构设计",
        "5 模块职责设计",
        "6 前后端交互设计",
        "7 数据存储设计",
        "8 安全设计",
        "9 部署运行设计",
    ])

    doc.add_heading("1 设计目标", 1)
    add_bullets(doc, [
        "建立统一的智能运维入口，使用户能够通过自然语言完成知识查询、数据分析和故障排查。",
        "将大模型能力限制在受控链路中，所有 SQL、脚本和文件路径操作都经过后端安全校验。",
        "通过日志、知识库和案例库形成故障处理闭环，使历史经验能够被后续问题复用。",
        "保持前后端模块清晰，便于后续扩展新数据源、新模型 Provider 和新诊断工具。",
    ])

    doc.add_heading("2 系统总体结构", 1)
    doc.add_paragraph("系统由 Vue 前端、FastAPI 后端、模型服务、向量存储、SQLite 配置库、文件型知识库、日志目录和外部业务数据源组成。前端负责页面交互和状态展示，后端负责接口校验、任务编排、模型调用、数据查询、安全控制和结果融合。")
    add_table(
        doc,
        ["组成部分", "技术", "说明"],
        [
            ["前端应用", "Vue 3、TypeScript、Vite、Element Plus、Pinia", "提供聊天、知识库、数据源、日志案例、诊断、索引和模型配置页面。"],
            ["后端服务", "FastAPI、Uvicorn", "提供 REST API 和 SSE 流式聊天接口。"],
            ["核心编排", "Orchestrator、IntentClassifier、TaskRouter、ResponseFusion", "完成意图识别、路由分发和回答融合。"],
            ["模型能力", "DeepSeek/OpenAI compatible、DashScope、BGE Embedding", "提供聊天生成、意图分类和向量化能力。"],
            ["向量存储", "Milvus Lite", "保存知识、日志和故障案例向量索引。"],
            ["配置存储", "SQLite、SQLAlchemy、Fernet", "保存数据源和 LLM Provider 配置。"],
            ["文件存储", "data/knowledge、data/uploads/logs、logs、scripts", "保存知识文件、上传日志、运行日志和诊断脚本。"],
        ],
        [1.2, 1.75, 3.55],
    )

    doc.add_heading("3 后端分层设计", 1)
    add_table(
        doc,
        ["层次", "主要文件/类", "职责"],
        [
            ["API 路由层", "ops_agent/api/routes/*.py", "接收 HTTP 请求，定义请求模型，调用服务层，返回 JSON 或 SSE。"],
            ["核心编排层", "Orchestrator、TaskRouter、ResponseFusion", "负责意图识别后的任务调度和多源证据融合。"],
            ["意图识别层", "IntentClassifier、IntentType", "通过规则和 LLM 判断 knowledge_query、data_analysis、fault_troubleshooting。"],
            ["RAG 层", "KnowledgeBase、Retriever、VectorStore", "加载知识文档，生成向量，执行相似检索。"],
            ["Text2SQL 层", "Text2SQLGenerator、SchemaManager、SQLValidator", "获取 schema，生成 SQL，执行安全校验。"],
            ["日志分析层", "LogUploadService、LogIndexer", "保存日志、发现日志、脱敏预览、建立日志索引。"],
            ["案例记忆层", "IncidentCaseMemory", "保存故障案例，按症状和 token 相似度复用案例。"],
            ["诊断工具层", "DiagnosticService、ScriptExecutor", "管理白名单脚本并限制执行。"],
            ["配置服务层", "config_service、llm_factory", "管理数据源、模型 Provider 和动态模型客户端。"],
        ],
        [1.15, 2.4, 3.0],
    )

    doc.add_heading("4 前端结构设计", 1)
    add_table(
        doc,
        ["页面", "路由", "调用接口", "展示内容"],
        [
            ["ChatView", "/", "/api/chat、/api/chat/stream、/api/uploads/logs", "聊天消息、流式回答、日志附件、来源、SQL、诊断信息。"],
            ["DataSourceView", "/datasources", "/api/config/datasources/*", "数据源列表、表单、连接测试、激活状态。"],
            ["KnowledgeView", "/knowledge", "/api/knowledge/*", "目录树、文件列表、文件预览、上传和重建索引。"],
            ["LogsCasesView", "/logs-cases", "/api/logs/*、/api/incidents/*", "日志列表、案例列表、分类、状态、详情抽屉。"],
            ["DiagnosticsView", "/diagnostics", "/api/diagnostics/*", "approved/pending 脚本列表、预览、启用、执行结果。"],
            ["IndexManagementView", "/indexes", "/api/indexes/*", "collection 状态、重建按钮、清理操作。"],
            ["LLMConfigView", "/llm", "/api/config/llm/*", "模型 Provider 列表、配置表单、测试和主力标记。"],
        ],
        [1.15, 1.1, 2.15, 2.1],
    )

    doc.add_heading("5 模块职责设计", 1)
    add_table(
        doc,
        ["模块", "主要职责", "关联接口", "关联数据"],
        [
            ["智能对话模块", "统一接收自然语言问题并展示答案。", "/api/chat、/api/chat/stream", "ChatRequest、ChatResponse、attachments。"],
            ["知识库模块", "管理知识文档、目录和知识向量索引。", "/api/knowledge/files、/tree、/upload、/reindex", "data/knowledge、ops_knowledge。"],
            ["数据源模块", "管理 MySQL、ClickHouse、Excel/CSV 数据源。", "/api/config/datasources/*", "datasource_configs。"],
            ["Text2SQL 模块", "自然语言到 SQL，并执行只读查询。", "/api/chat 间接调用", "业务数据库 schema、rows。"],
            ["日志模块", "上传、发现、预览、脱敏、分析日志。", "/api/uploads/logs、/api/logs/*", "log_metadata、ops_logs。"],
            ["故障案例模块", "保存和复用故障处理经验。", "/api/incidents/*", "incident_cases、ops_incident_cases。"],
            ["诊断模块", "管理和执行白名单脚本。", "/api/diagnostics/*", "scripts/approved、pending、disabled。"],
            ["索引模块", "维护知识、日志、案例索引。", "/api/indexes/*", "Milvus Lite collections。"],
            ["模型配置模块", "管理 LLM Provider 和主力模型。", "/api/config/llm/*", "llm_provider_configs。"],
        ],
        [1.15, 2.25, 1.95, 1.75],
    )

    doc.add_heading("6 前后端交互设计", 1)
    add_numbered(doc, [
        "用户在前端页面进行操作，例如发送问题、上传日志、创建数据源或运行脚本。",
        "前端 api 层使用 Axios 或 EventSource/fetch 调用后端接口。",
        "后端路由层完成参数校验，非法输入直接返回 400 或 404。",
        "服务层执行具体业务逻辑，包括文件保存、数据库查询、索引重建、脚本执行等。",
        "后端返回 JSON；流式聊天接口返回 text/event-stream。",
        "前端根据返回字段更新列表、详情、消息内容或错误提示。",
    ])

    doc.add_heading("7 数据存储设计", 1)
    add_table(
        doc,
        ["数据类型", "存储位置", "用途"],
        [
            ["配置数据", "data/app_config.db", "保存数据源和 LLM Provider。"],
            ["故障案例", "data/incident_cases.db", "保存 query、answer、symptoms、root_cause、solution、status、category。"],
            ["知识文件", "data/knowledge", "保存 md/txt 知识文档。"],
            ["上传日志", "data/uploads/logs", "保存上传文件和 metadata JSON。"],
            ["运行日志", "logs", "系统运行日志，可被日志目录服务发现。"],
            ["向量索引", "data/vectors/milvus.db", "Milvus Lite collection 文件。"],
            ["诊断脚本", "scripts/approved、pending、disabled", "保存可执行脚本、待审核脚本和停用脚本。"],
        ],
        [1.25, 2.2, 3.05],
    )

    doc.add_heading("8 安全设计", 1)
    add_table(
        doc,
        ["安全点", "设计"],
        [
            ["SQL 安全", "SQLValidator 禁止写操作，要求 SELECT 和 LIMIT，LIMIT 不超过 200，并检测常见注入片段。"],
            ["脚本安全", "DiagnosticService 只允许 check_*.sh 和 check_*.py，执行时必须位于 approved 目录。"],
            ["路径安全", "知识、日志和脚本服务使用 resolve 后的路径校验，确保目标在根目录内。"],
            ["密钥安全", "数据源密码和 LLM API Key 使用 Fernet 加密保存在 SQLite。"],
            ["日志脱敏", "预览和上下文提取时替换 password、token、secret、api_key 和 Bearer token。"],
            ["访问控制", "非 debug 模式启用 APIKeyMiddleware。"],
        ],
        [1.5, 5.0],
    )

    doc.add_heading("9 部署运行设计", 1)
    add_table(
        doc,
        ["事项", "说明"],
        [
            ["后端启动", "uvicorn ops_agent.api.main:app --reload --port 8080。"],
            ["前端开发", "cd frontend && npm run dev，Vite 将 /api 代理到 localhost:8080。"],
            ["前端构建", "cd frontend && npm run build，输出到 ops_agent/api/static/dist。"],
            ["静态托管", "FastAPI 挂载 /assets，并通过 SPA fallback 返回 index.html。"],
            ["运行产物", "data/uploads、data/vectors、data/app_config.db、data/incident_cases.db 不应提交到 git。"],
        ],
        [1.4, 5.1],
    )

    blacken_runs(doc)
    doc.save(OUT / "OpsAgent_中期验收_概要设计.docx")


API_ROWS = [
    ["POST", "/api/chat", "非流式聊天", "Body: query:string; history:list[dict]=[]; datasource_id:string|null; attachments:list[dict]=[]", "answer:string; intent:string; sources:list[dict]; sql:string; diagnostics:dict"],
    ["POST", "/api/chat/stream", "SSE 流式聊天", "Body 同 /api/chat", "SSE intent:{type}; token:{text}; done:{intent,sources,sql,diagnostics}; error:{error}"],
    ["GET", "/api/config/datasources", "数据源列表", "无", "list[{id,name,type,is_active,config,created_at,updated_at}]"],
    ["POST", "/api/config/datasources", "创建数据源", "Body: name; type=mysql|clickhouse|excel_csv; config; is_active", "DataSource 对象"],
    ["POST", "/api/config/datasources/upload-file", "上传 Excel/CSV 数据文件", "Form file，支持 .csv/.xlsx/.xls，最大 100MB", "upload_id,file_path,original_filename,file_type,size_bytes,sheet_names"],
    ["PUT", "/api/config/datasources/{ds_id}", "更新数据源", "Path ds_id; Body: name/type/config/is_active 可选", "DataSource 对象；不存在返回 404"],
    ["DELETE", "/api/config/datasources/{ds_id}", "删除数据源", "Path ds_id", "{ok:true}; 不存在返回 404"],
    ["POST", "/api/config/datasources/{ds_id}/activate", "激活数据源", "Path ds_id", "{ok:true}; 同时取消其他 active"],
    ["POST", "/api/config/datasources/{ds_id}/test", "测试已保存数据源", "Path ds_id", "{ok:boolean,message:string,latency_ms:number}"],
    ["POST", "/api/config/datasources/tables", "查看未保存数据源表", "Body: type, config", "{tables:list[string]}"],
    ["GET", "/api/config/datasources/{ds_id}/tables", "查看已保存数据源表", "Path ds_id", "{tables:list[string]}"],
    ["POST", "/api/config/datasources/test", "测试未保存数据源", "Body: DataSourceCreate", "{ok:boolean,message:string,latency_ms:number}"],
    ["GET", "/api/config/llm", "模型 Provider 列表", "无", "list[{id,name,provider_type,base_url,model,temperature,max_tokens,is_primary,api_key,created_at,updated_at}]"],
    ["POST", "/api/config/llm", "创建模型 Provider", "Body: name,provider_type,api_key,base_url,model,temperature,max_tokens,is_primary", "LLMProvider 对象"],
    ["PUT", "/api/config/llm/{prov_id}", "更新模型 Provider", "Path prov_id; Body 字段可选", "LLMProvider 对象；不存在返回 404"],
    ["DELETE", "/api/config/llm/{prov_id}", "删除模型 Provider", "Path prov_id", "{ok:true}; 不存在返回 404"],
    ["POST", "/api/config/llm/{prov_id}/primary", "设置主力模型", "Path prov_id", "{ok:true}; 同时取消其他 primary"],
    ["POST", "/api/config/llm/{prov_id}/test", "测试已保存 Provider", "Path prov_id; Body: message:string", "{response:string,latency_ms:number,error?:string}"],
    ["POST", "/api/config/llm/test", "测试未保存 Provider", "Body: name,provider_type,api_key,base_url,model,temperature,max_tokens,message", "{response:string,latency_ms:number,error?:string}"],
    ["GET", "/api/knowledge/files", "知识文件列表", "无", "list[{file_id,filename,relative_path,size,updated_at,indexed}]"],
    ["GET", "/api/knowledge/tree", "知识库目录树", "无", "tree nodes: name,relative_path,children,files"],
    ["POST", "/api/knowledge/folders", "创建目录", "Body: path:string", "{name,relative_path}"],
    ["DELETE", "/api/knowledge/folders", "删除目录", "Query: path:string; recursive:boolean=false", "{deleted:boolean}"],
    ["PUT", "/api/knowledge/folders/rename", "重命名目录", "Body: path,new_name", "{name,relative_path}"],
    ["POST", "/api/knowledge/upload", "上传知识文件", "Query: filename,folder; Body: raw bytes", "{file_id,filename,relative_path,size,updated_at,indexed}"],
    ["GET", "/api/knowledge/files/{file_id}", "获取知识文件", "Path file_id", "metadata + content"],
    ["DELETE", "/api/knowledge/files/{file_id}", "删除知识文件", "Path file_id", "{deleted:boolean}"],
    ["POST", "/api/knowledge/reindex", "重建知识库索引", "无", "{status,collection,count,indexed_at}"],
    ["POST", "/api/uploads/logs", "上传日志附件", "Query: filename,category; Body: raw bytes", "LogMetadata"],
    ["GET", "/api/logs", "日志列表", "Query: query,category,source,severity", "list[LogMetadata]"],
    ["GET", "/api/logs/categories", "日志分类统计", "无", "list[{name,count,error_count,warning_count,pinned,user_defined}]"],
    ["POST", "/api/logs/categories", "创建日志分类", "Body: name", "{name,pinned,user_defined}"],
    ["PUT", "/api/logs/categories/rename", "重命名日志分类", "Body: old_name,new_name", "category object"],
    ["PUT", "/api/logs/categories/pin", "置顶日志分类", "Body: name,pinned", "category object"],
    ["DELETE", "/api/logs/categories", "删除日志分类", "Body: name", "{deleted:boolean}"],
    ["GET", "/api/logs/{file_id}", "日志 metadata", "Path file_id", "LogMetadata"],
    ["GET", "/api/logs/{file_id}/preview", "日志脱敏预览", "Path file_id", "LogMetadata + content"],
    ["PUT", "/api/logs/{file_id}/category", "更新日志分类", "Path file_id; Body: category", "{updated:boolean}"],
    ["DELETE", "/api/logs/{file_id}", "删除日志", "Path file_id", "{deleted:boolean}"],
    ["GET", "/api/incidents", "故障案例列表", "Query: status,query,category,symptom", "list[IncidentCase]"],
    ["GET", "/api/incidents/categories", "案例分类统计", "无", "list[{name,count,pinned,user_defined}]"],
    ["POST", "/api/incidents/categories", "创建案例分类", "Body: name", "category object"],
    ["PUT", "/api/incidents/categories/rename", "重命名案例分类", "Body: old_name,new_name", "category object"],
    ["PUT", "/api/incidents/categories/pin", "置顶案例分类", "Body: name,pinned", "category object"],
    ["DELETE", "/api/incidents/categories", "删除案例分类", "Body: name", "{deleted:boolean}"],
    ["GET", "/api/incidents/{case_id}", "案例详情", "Path case_id", "IncidentCase"],
    ["PUT", "/api/incidents/{case_id}/status", "更新案例状态", "Path case_id; Body: status", "{updated:boolean}"],
    ["PUT", "/api/incidents/{case_id}/category", "更新案例分类", "Path case_id; Body: category", "{updated:boolean}"],
    ["DELETE", "/api/incidents/{case_id}", "删除案例", "Path case_id", "{deleted:boolean}"],
    ["GET", "/api/diagnostics/scripts", "approved 脚本列表", "无", "list[{name,size,description,timeout}]"],
    ["GET", "/api/diagnostics/pending", "pending 脚本列表", "无", "list[{name,size,description,timeout}]"],
    ["POST", "/api/diagnostics/upload", "上传待启用脚本", "Query: filename; Body: raw bytes", "{name,size,description,timeout,status}"],
    ["GET", "/api/diagnostics/scripts/{script_name}/preview", "预览脚本", "Path script_name; Query: status=approved", "metadata + content"],
    ["POST", "/api/diagnostics/scripts/{script_name}/enable", "启用脚本", "Path script_name", "{name,status:'enabled'}"],
    ["POST", "/api/diagnostics/scripts/{script_name}/disable", "停用脚本", "Path script_name", "{name,status:'disabled'}"],
    ["DELETE", "/api/diagnostics/scripts/{script_name}", "删除脚本", "Path script_name; Query: status=pending", "{deleted:boolean}"],
    ["POST", "/api/diagnostics/scripts/{script_name}/run", "执行脚本", "Path script_name; Body: args:list[string]", "{stdout,stderr,exit_code}"],
    ["GET", "/api/indexes/status", "索引状态", "无", "{milvus_db_path,knowledge_dir,log_dir,collections}"],
    ["POST", "/api/indexes/knowledge/rebuild", "重建知识索引", "无", "{status,collection,count,indexed_at}"],
    ["POST", "/api/indexes/logs/rebuild", "重建日志索引", "Body: path:string|null", "{status,collection,target,targets,count}"],
    ["POST", "/api/indexes/cases/rebuild", "重建案例索引", "无", "{status,collection,count}"],
    ["POST", "/api/indexes/{collection}/clear", "清空 collection", "Path collection", "{status:'cleared',collection}"],
]


def build_detail_doc():
    draw_er_diagram()
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "OpsAgent 中期验收 - 详细设计")
    add_cover(doc, "OpsAgent 智能运维助手", "详细设计说明书")
    add_toc(doc, [
        "1 详细设计说明",
        "2 后端接口设计",
        "3 请求与响应对象",
        "4 数据库与 E-R 设计",
        "5 核心类与模块设计",
        "6 关键业务流程",
        "7 安全设计",
        "8 测试设计",
    ])

    doc.add_heading("1 详细设计说明", 1)
    doc.add_paragraph("详细设计文档说明 OpsAgent 后端接口、请求参数、返回参数、数据库结构、核心类、模块交互、业务流程和安全控制。接口内容依据当前 FastAPI 路由和服务实现整理。")

    doc.add_heading("2 后端接口设计", 1)
    add_table(doc, ["方法", "路径", "功能", "请求参数", "返回参数"], API_ROWS, [0.55, 1.55, 1.25, 2.3, 2.35])

    doc.add_heading("3 请求与响应对象", 1)
    doc.add_heading("3.1 ChatRequest / ChatResponse", 2)
    add_table(
        doc,
        ["对象", "字段", "类型", "说明"],
        [
            ["ChatRequest", "query", "string", "用户输入的自然语言问题，必填。"],
            ["ChatRequest", "history", "list[dict]", "历史对话数组，默认空数组。"],
            ["ChatRequest", "datasource_id", "string|null", "指定数据源 id，可为空。"],
            ["ChatRequest", "attachments", "list[dict]", "附件数组；日志附件 type 为 log，id 为 file_id。"],
            ["ChatResponse", "answer", "string", "最终回答文本。"],
            ["ChatResponse", "intent", "string", "knowledge_query、data_analysis、fault_troubleshooting 或 error。"],
            ["ChatResponse", "sources", "list[dict]", "知识来源、日志来源或案例来源。"],
            ["ChatResponse", "sql", "string", "数据分析场景生成并执行的 SQL。"],
            ["ChatResponse", "diagnostics", "dict", "故障排查元数据，包括脚本输出、案例命中和附件信息。"],
        ],
        [1.1, 1.25, 1.25, 3.0],
    )
    doc.add_heading("3.2 DataSourceCreate", 2)
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ["name", "string", "数据源名称。"],
            ["type", "string", "mysql、clickhouse 或 excel_csv。"],
            ["config.host", "string|null", "数据库主机地址。"],
            ["config.port", "int|null", "数据库端口，MySQL 默认 3306，ClickHouse 默认 8123。"],
            ["config.user", "string|null", "数据库用户名。"],
            ["config.password", "string|null", "数据库密码，保存时加密。"],
            ["config.database", "string|null", "数据库名。"],
            ["config.charset", "string|null", "MySQL 字符集，默认 utf8mb4。"],
            ["config.file_path", "string|null", "Excel/CSV 文件路径。"],
            ["config.files", "list[ExcelCSVFileSchema]|null", "Excel/CSV 多文件配置，最多 5 个文件。"],
            ["is_active", "boolean", "是否设为活跃数据源。"],
        ],
        [1.75, 1.6, 3.15],
    )
    doc.add_heading("3.3 LLMProviderCreate", 2)
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ["name", "string", "Provider 名称。"],
            ["provider_type", "string", "openai_compatible 或 dashscope。"],
            ["api_key", "string", "模型 API Key，保存时加密。"],
            ["base_url", "string", "OpenAI 兼容接口基础地址。"],
            ["model", "string", "模型名称。"],
            ["temperature", "float", "采样温度，默认 0.1。"],
            ["max_tokens", "int", "最大输出 token，默认 4096。"],
            ["is_primary", "boolean", "是否设为主力模型。"],
        ],
        [1.6, 1.4, 3.5],
    )

    doc.add_heading("4 数据库与 E-R 设计", 1)
    add_er(doc)
    add_table(
        doc,
        ["关系", "说明"],
        [
            ["datasource_configs - incident_cases", "数据分析和故障排查可使用业务数据源结果作为诊断证据。"],
            ["llm_provider_configs - incident_cases", "故障案例的根因、方案和总结由当前主力模型参与生成。"],
            ["knowledge_file - Milvus collections", "知识文件分块后写入 ops_knowledge collection，供知识问答和故障排查检索。"],
            ["log_metadata - Milvus collections", "上传日志和本地日志可写入 ops_logs collection，供日志相似检索。"],
            ["incident_cases - Milvus collections", "案例文本写入 ops_incident_cases collection，供历史案例复用。"],
            ["category_registry - log_metadata", "日志分类信息用于日志列表筛选和分类统计。"],
            ["category_registry - incident_cases", "案例分类信息用于案例列表筛选和复盘管理。"],
        ],
        [2.1, 4.4],
    )
    doc.add_heading("4.1 datasource_configs 表", 2)
    add_table(
        doc,
        ["字段名", "类型", "主键", "可空", "默认值", "描述"],
        [
            ["id", "String(36)", "是", "否", "无", "数据源 id，UUID 字符串。"],
            ["name", "String(128)", "否", "否", "无", "数据源显示名称。"],
            ["type", "String(32)", "否", "否", "无", "mysql、clickhouse 或 excel_csv。"],
            ["config_json", "Text", "否", "否", "无", "数据源配置 JSON，密码字段加密保存。"],
            ["is_active", "Boolean", "否", "是", "False", "是否为当前活跃数据源。"],
            ["created_at", "DateTime", "否", "是", "当前 UTC 时间", "创建时间。"],
            ["updated_at", "DateTime", "否", "是", "当前 UTC 时间", "更新时间。"],
        ],
        [1.1, 1.05, 0.55, 0.55, 1.3, 2.25],
    )
    doc.add_heading("4.2 llm_provider_configs 表", 2)
    add_table(
        doc,
        ["字段名", "类型", "主键", "可空", "默认值", "描述"],
        [
            ["id", "String(36)", "是", "否", "无", "Provider id，UUID 字符串。"],
            ["name", "String(128)", "否", "否", "无", "Provider 名称。"],
            ["provider_type", "String(32)", "否", "否", "无", "openai_compatible 或 dashscope。"],
            ["api_key_encrypted", "Text", "否", "否", "无", "加密后的 API Key。"],
            ["base_url", "String(512)", "否", "否", "无", "接口基础地址。"],
            ["model", "String(128)", "否", "否", "无", "模型名称。"],
            ["temperature", "Float", "否", "是", "0.1", "模型采样温度。"],
            ["max_tokens", "Integer", "否", "是", "4096", "最大输出 token。"],
            ["is_primary", "Boolean", "否", "是", "False", "是否为主力模型。"],
            ["created_at", "DateTime", "否", "是", "当前 UTC 时间", "创建时间。"],
            ["updated_at", "DateTime", "否", "是", "当前 UTC 时间", "更新时间。"],
        ],
        [1.1, 1.05, 0.55, 0.55, 1.2, 2.35],
    )
    doc.add_heading("4.3 incident_cases 表", 2)
    add_table(
        doc,
        ["字段名", "类型", "主键", "可空", "默认值", "描述"],
        [
            ["case_id", "TEXT", "是", "否", "无", "故障案例编号，case_<uuid>。"],
            ["query", "TEXT", "否", "否", "无", "用户原始问题。"],
            ["answer", "TEXT", "否", "否", "无", "系统生成的诊断回答。"],
            ["symptoms", "TEXT(JSON)", "否", "否", "[]", "故障症状数组。"],
            ["root_cause", "TEXT", "否", "否", "空字符串", "根因摘要。"],
            ["solution", "TEXT", "否", "否", "空字符串", "解决方案摘要。"],
            ["evidence", "TEXT(JSON)", "否", "否", "[]", "日志、知识、脚本等证据。"],
            ["status", "TEXT", "否", "否", "pending", "pending、resolved、auto_saved 等。"],
            ["category", "TEXT", "否", "否", "空字符串", "案例分类。"],
            ["tokens", "TEXT(JSON)", "否", "否", "[]", "相似案例匹配 token。"],
            ["created_at", "TEXT", "否", "否", "无", "创建时间。"],
            ["updated_at", "TEXT", "否", "否", "无", "更新时间。"],
        ],
        [1.05, 1.0, 0.5, 0.5, 1.15, 2.4],
    )
    doc.add_heading("4.4 log_metadata JSON", 2)
    add_table(
        doc,
        ["字段名", "类型", "描述"],
        [
            ["file_id", "string", "日志编号。上传日志为 log_<uuid>，本地发现日志为 local_<sha1>。"],
            ["filename", "string", "安全文件名。"],
            ["size", "int", "文件字节数。"],
            ["category", "string", "日志分类。"],
            ["source", "string", "uploaded、runtime、seed 等来源。"],
            ["severity", "string", "根据分析结果推断的严重级别。"],
            ["tags", "list[string]", "关键模式标签。"],
            ["stored_path", "string", "实际文件路径，预览前校验目录边界。"],
            ["uploaded_at", "string", "上传时间，本地发现日志可为空。"],
            ["updated_at", "string", "metadata 更新时间。"],
            ["mtime", "float", "文件修改时间戳。"],
            ["analysis.error_count", "int", "错误行数量。"],
            ["analysis.warning_count", "int", "警告行数量。"],
            ["analysis.patterns", "list[string]", "Connection refused、Permission denied、OOM 等关键模式。"],
            ["analysis.summary", "string", "脱敏后的关键片段摘要。"],
        ],
        [1.75, 1.45, 3.3],
    )
    doc.add_heading("4.5 knowledge_file metadata", 2)
    add_table(
        doc,
        ["字段名", "类型", "描述"],
        [
            ["file_id", "string", "relative_path 的 base64 urlsafe 编码。"],
            ["filename", "string", "文件名。"],
            ["relative_path", "string", "相对 data/knowledge 的路径。"],
            ["size", "int", "文件字节数。"],
            ["updated_at", "string", "文件修改时间。"],
            ["indexed", "boolean", "文件修改时间是否早于最近索引时间。"],
            ["content", "string", "获取详情接口返回的文件内容。"],
        ],
        [1.7, 1.4, 3.4],
    )
    doc.add_heading("4.6 Milvus collections", 2)
    add_table(
        doc,
        ["Collection", "内容", "来源"],
        [
            ["ops_knowledge", "知识文档分块向量。", "KnowledgeService.rebuild_index。"],
            ["ops_logs", "日志分块向量。", "IndexService.rebuild_logs / LogIndexer。"],
            ["ops_incident_cases", "故障案例文本向量。", "IndexService.rebuild_cases。"],
        ],
        [1.55, 2.55, 2.4],
    )

    doc.add_heading("5 核心类与模块设计", 1)
    add_table(
        doc,
        ["类/模块", "主要方法", "职责"],
        [
            ["Orchestrator", "process、process_stream", "聊天总入口，合并日志附件，识别意图，调用任务处理器，返回普通结果或 SSE 事件。"],
            ["IntentClassifier", "classify、_rule_classify、_llm_classify", "规则和 LLM 两级意图识别，抽取 IP、主机名、服务名和端口。"],
            ["TaskRouter", "register、route", "维护 IntentType 到处理函数的映射。"],
            ["ResponseFusion", "fuse_for_knowledge、fuse_for_data、fuse_for_troubleshooting", "将知识、数据、日志、脚本和案例证据整理为回答。"],
            ["Text2SQLGenerator", "generate", "基于 schema 和用户问题生成 SQL。"],
            ["SQLValidator", "validate", "执行 SQL 安全校验。"],
            ["LogUploadService", "save_log_file、list_logs、preview_log、resolve_mentioned_logs", "管理上传日志、本地日志发现、脱敏预览和附件上下文。"],
            ["IncidentCaseMemory", "save_case、find_similar、list_cases", "保存和检索故障案例。"],
            ["DiagnosticService", "upload_script、enable_script、run_script", "管理 pending/approved/disabled 脚本并执行白名单脚本。"],
            ["IndexService", "status、rebuild_knowledge、rebuild_logs、rebuild_cases、clear_collection", "管理 Milvus collection 状态和重建。"],
        ],
        [1.35, 2.15, 3.0],
    )

    doc.add_heading("6 关键业务流程", 1)
    doc.add_heading("6.1 聊天请求处理流程", 2)
    add_numbered(doc, [
        "前端提交 ChatRequest。",
        "chat.py 调用 Orchestrator.process 或 process_stream。",
        "Orchestrator 解析用户文本中提到的日志文件名，并与显式 attachments 合并。",
        "IntentClassifier 判断意图；如果附件中存在 type=log，则强制切换为 fault_troubleshooting。",
        "TaskRouter 根据意图调用知识问答、数据分析或故障排查处理器。",
        "处理器返回 answer、sources、sql、diagnostics。",
        "普通接口返回 ChatResponse，流式接口返回 SSE 事件。",
    ])
    doc.add_heading("6.2 Text2SQL 流程", 2)
    add_numbered(doc, [
        "获取活跃数据源，或使用 datasource_id 指定的数据源。",
        "SchemaManager 读取表结构并提供给 Text2SQLGenerator。",
        "LLM 生成候选 SQL。",
        "SQLValidator 校验危险关键字、SELECT、LIMIT 和注入模式。",
        "数据源执行 SQL，返回 rows。",
        "ResponseFusion 根据 rows 生成自然语言总结。",
    ])
    doc.add_heading("6.3 故障排查流程", 2)
    add_numbered(doc, [
        "收集 query、实体、上传日志上下文和文本中提到的日志。",
        "IncidentCaseMemory.find_similar 先查找 resolved/auto_saved 相似案例。",
        "命中案例时直接返回历史根因和解决方案。",
        "未命中时并行获取知识库上下文、日志检索结果和诊断脚本输出。",
        "ResponseFusion 生成诊断回答。",
        "系统把诊断结果保存为故障案例，用于后续复用。",
    ])
    doc.add_heading("6.4 诊断脚本流程", 2)
    add_numbered(doc, [
        "用户上传脚本，文件名必须匹配 check_*.sh 或 check_*.py。",
        "脚本保存到 scripts/pending，默认不可执行。",
        "管理员预览后启用，脚本移动到 scripts/approved。",
        "执行接口只接收 approved 目录中的脚本名。",
        "执行时设置 cwd 为 approved 目录，超时 30 秒，输出截断 5000 字符。",
    ])

    doc.add_heading("7 安全设计", 1)
    add_table(
        doc,
        ["安全项", "具体规则", "实现位置"],
        [
            ["SQL 安全", "禁止 DROP/DELETE/INSERT/UPDATE/ALTER/TRUNCATE/CREATE/EXEC/GRANT/REVOKE/UNION；只允许 SELECT；必须 LIMIT；LIMIT <= 200。", "SQLValidator"],
            ["日志脱敏", "替换 password/passwd/pwd、token、secret、api_key、Authorization Bearer。", "LogUploadService._redact"],
            ["路径安全", "知识文件、日志文件、脚本文件 resolve 后必须位于对应根目录。", "KnowledgeService、LogUploadService、DiagnosticService"],
            ["脚本安全", "上传脚本限制名称和大小，启用后才可执行，只执行 approved 目录脚本。", "DiagnosticService"],
            ["配置安全", "数据源密码和 LLM API Key 使用 Fernet 加密。", "config_service"],
            ["接口安全", "debug=false 时启用 APIKeyMiddleware。", "ops_agent/api/main.py"],
        ],
        [1.1, 3.75, 1.65],
    )

    doc.add_heading("8 测试设计", 1)
    add_table(
        doc,
        ["测试文件", "测试重点"],
        [
            ["tests/test_intent.py", "意图分类、日志文件名识别、数据查询和知识查询规则。"],
            ["tests/test_text2sql.py", "SQLValidator 的危险关键字、LIMIT 和注入模式校验。"],
            ["tests/test_log_upload_service.py", "日志上传、脱敏、本地日志发现、附件解析。"],
            ["tests/test_chat_attachments_contract.py", "日志附件强制进入故障排查。"],
            ["tests/test_incident_case_memory.py", "案例保存、相似匹配、状态和分类管理。"],
            ["tests/test_management_services.py", "知识库、诊断、索引等管理服务。"],
            ["tests/test_orchestrator_contract.py", "编排器返回结构和契约。"],
            ["frontend 类型检查", "cd frontend && npx vue-tsc -p tsconfig.app.json --noEmit。"],
        ],
        [2.2, 4.3],
    )

    blacken_runs(doc)
    doc.save(OUT / "OpsAgent_中期验收_详细设计.docx")


def build_all():
    OUT.mkdir(parents=True, exist_ok=True)
    build_requirements_doc()
    build_outline_doc()
    build_detail_doc()


if __name__ == "__main__":
    build_all()
    for path in sorted(OUT.glob("*.docx")):
        print(path)
